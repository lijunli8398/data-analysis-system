#!/usr/bin/env python3
"""
Word报告生成器 - 生成真实的数据分析报告
"""

import os
import sys
import json
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 添加 skill 路径
SKILL_PATH = Path(__file__).parent.parent.parent / "skills" / "data-analysis"
if SKILL_PATH.exists():
    sys.path.insert(0, str(SKILL_PATH))

from modules.llm_client import llm_client


class ReportGenerator:
    """Word报告生成器 - 生成真实分析报告"""
    
    def generate(self, data_path: str, title: str, prompt: str, output_dir: Path) -> dict:
        """生成完整分析报告
        
        Args:
            data_path: 数据文件路径
            title: 报告标题
            prompt: 用户分析提示词（分析需求、报告结构等）
            output_dir: 输出目录
        """
        
        report_id = datetime.now().strftime("%Y%m%d%H%M%S")
        
        # 创建目录
        charts_dir = output_dir / "charts"
        charts_dir.mkdir(parents=True, exist_ok=True)
        results_dir = output_dir / "results"
        results_dir.mkdir(parents=True, exist_ok=True)
        reports_dir = output_dir / "reports"
        reports_dir.mkdir(parents=True, exist_ok=True)
        
        try:
            # ========== 第一步：加载数据 ==========
            import pandas as pd
            import numpy as np
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import seaborn as sns
            
            # 中文字体配置
            plt.rcParams['font.sans-serif'] = ['Noto Serif CJK SC', 'WenQuanYi Zen Hei', 'SimHei']
            plt.rcParams['axes.unicode_minus'] = False
            
            # 加载数据
            if data_path.endswith('.csv'):
                df = pd.read_csv(data_path)
            elif data_path.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(data_path)
            else:
                df = pd.read_csv(data_path)
            
            # ========== 第二步：执行数据分析（传入用户提示词）==========
            analysis_results = self._perform_analysis(df, charts_dir, user_prompt=prompt)
            
            # ========== 3. 使用智能分析引擎 ==========
            # 导入智能分析引擎
            try:
                from modules.smart_analysis_engine import SmartAnalysisEngine
                engine = SmartAnalysisEngine()
                smart_result = engine.analyze(df)
                
                # 如果智能分析质量更高，使用智能分析结果
                if smart_result.get('quality_score', 0) > 0.5:
                    analysis_results['smart_analysis'] = smart_result
                    analysis_results['data_profile'] = smart_result['profile']
            except Exception as e:
                print(f"智能分析引擎调用失败: {e}")
            
            # ========== 4. 生成Word报告（传入用户提示词）==========
            docx_path = reports_dir / f"{report_id}.docx"
            self._generate_word_report(title, analysis_results, docx_path, user_prompt=prompt)
            
            # 保存分析结果JSON
            result_json_path = results_dir / f"{report_id}.json"
            with open(result_json_path, 'w', encoding='utf-8') as f:
                json.dump(analysis_results, f, ensure_ascii=False, indent=2, default=str)
            
            return {
                'report_id': report_id,
                'docx_path': str(docx_path),
                'result_path': str(result_json_path),
                'charts': analysis_results.get('chart_paths', [])
            }
            
        except Exception as e:
            raise Exception(f"报告生成失败: {str(e)}")
    
    def _perform_analysis(self, df, charts_dir: Path, user_prompt: str = "") -> dict:
        """执行完整的数据分析
        
        Args:
            df: 数据DataFrame
            charts_dir: 图表输出目录
            user_prompt: 用户分析提示词，用于定制分析内容和报告结构
        """
        
        import numpy as np
        import matplotlib.pyplot as plt
        import seaborn as sns
        
        results = {
            'basic_info': {},
            'data_quality': {},
            'numeric_analysis': {},
            'categorical_analysis': {},
            'chart_paths': [],
            'user_prompt': user_prompt  # 保存用户提示词
        }
        
        # ========== 1. 基本信息 ==========
        results['basic_info'] = {
            'total_rows': int(len(df)),
            'total_columns': int(len(df.columns)),
            'memory_usage': f"{df.memory_usage(deep=True).sum() / 1024 / 1024:.2f} MB",
            'duplicate_rows': int(df.duplicated().sum())
        }
        
        # ========== 2. 数据质量分析 ==========
        missing_info = {}
        for col in df.columns:
            missing_count = df[col].isnull().sum()
            missing_pct = (missing_count / len(df)) * 100
            if missing_count > 0:
                missing_info[col] = {
                    'missing_count': int(missing_count),
                    'missing_percentage': f"{missing_pct:.2f}%"
                }
        
        results['data_quality'] = {
            'columns_with_missing': len(missing_info),
            'missing_details': missing_info,
            'completeness_rate': f"{((len(df.columns) - len(missing_info)) / len(df.columns)) * 100:.1f}%"
        }
        
        # ========== 3. 数值型变量分析 ==========
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        
        if numeric_cols:
            numeric_stats = {}
            for col in numeric_cols:
                col_data = df[col].dropna()
                if len(col_data) > 0:
                    numeric_stats[col] = {
                        'count': int(len(col_data)),
                        'mean': float(col_data.mean()),
                        'std': float(col_data.std()),
                        'min': float(col_data.min()),
                        'max': float(col_data.max()),
                        'median': float(col_data.median()),
                        'q1': float(col_data.quantile(0.25)),
                        'q3': float(col_data.quantile(0.75)),
                        'skewness': float(col_data.skew()),
                        'kurtosis': float(col_data.kurtosis())
                    }
            
            results['numeric_analysis'] = {
                'numeric_columns': numeric_cols,
                'statistics': numeric_stats
            }
            
            # 生成数值变量分布图
            for i, col in enumerate(numeric_cols[:6]):
                try:
                    fig, axes = plt.subplots(1, 2, figsize=(12, 4))
                    
                    # 直方图
                    axes[0].hist(df[col].dropna(), bins=30, color='#667eea', edgecolor='white', alpha=0.8)
                    axes[0].set_title(f'{col} 分布', fontsize=12)
                    axes[0].set_xlabel(col)
                    axes[0].set_ylabel('频次')
                    axes[0].grid(alpha=0.3)
                    
                    # 箱线图
                    axes[1].boxplot(df[col].dropna(), vert=True, patch_artist=True,
                                   boxprops=dict(facecolor='#667eea', alpha=0.7))
                    axes[1].set_title(f'{col} 箱线图', fontsize=12)
                    axes[1].set_ylabel(col)
                    axes[1].grid(alpha=0.3)
                    
                    plt.tight_layout()
                    chart_path = str(charts_dir / f'numeric_{col}.png')
                    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
                    plt.close()
                    results['chart_paths'].append(chart_path)
                except Exception as e:
                    print(f"图表生成失败 {col}: {e}")
        
        # ========== 4. 分类型变量分析 ==========
        categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        
        if categorical_cols:
            cat_analysis = {}
            for col in categorical_cols:
                value_counts = df[col].value_counts()
                unique_count = df[col].nunique()
                
                cat_analysis[col] = {
                    'unique_values': int(unique_count),
                    'top_values': value_counts.head(10).to_dict(),
                    'top_percentage': {k: f"{v/len(df)*100:.1f}%" for k, v in value_counts.head(5).items()}
                }
            
            results['categorical_analysis'] = {
                'categorical_columns': categorical_cols,
                'analysis': cat_analysis
            }
            
            # 生成分类型变量柱状图
            for col in categorical_cols[:4]:
                try:
                    value_counts = df[col].value_counts().head(10)
                    
                    fig, ax = plt.subplots(figsize=(10, 5))
                    bars = ax.bar(range(len(value_counts)), value_counts.values, 
                                 color='#667eea', edgecolor='white')
                    ax.set_xticks(range(len(value_counts)))
                    ax.set_xticklabels([str(x)[:15] for x in value_counts.index], rotation=45, ha='right')
                    ax.set_title(f'{col} 分布 (Top 10)', fontsize=12)
                    ax.set_ylabel('数量')
                    ax.grid(alpha=0.3, axis='y')
                    
                    # 添加数值标签
                    for bar, val in zip(bars, value_counts.values):
                        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                               f'{val}', ha='center', va='bottom', fontsize=9)
                    
                    plt.tight_layout()
                    chart_path = str(charts_dir / f'categorical_{col}.png')
                    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
                    plt.close()
                    results['chart_paths'].append(chart_path)
                except Exception as e:
                    print(f"分类图表生成失败 {col}: {e}")
        
        # ========== 5. 相关性分析 ==========
        if len(numeric_cols) >= 2:
            try:
                corr_matrix = df[numeric_cols].corr()
                
                fig, ax = plt.subplots(figsize=(10, 8))
                sns.heatmap(corr_matrix, annot=True, fmt='.2f', cmap='RdBu_r', 
                           center=0, square=True, linewidths=0.5, ax=ax)
                ax.set_title('数值变量相关性热力图', fontsize=14)
                plt.tight_layout()
                chart_path = str(charts_dir / 'correlation_heatmap.png')
                plt.savefig(chart_path, dpi=150, bbox_inches='tight')
                plt.close()
                results['chart_paths'].append(chart_path)
                
                # 找出强相关变量对
                strong_correlations = []
                for i in range(len(numeric_cols)):
                    for j in range(i+1, len(numeric_cols)):
                        corr_val = corr_matrix.iloc[i, j]
                        if abs(corr_val) > 0.5:
                            strong_correlations.append({
                                'var1': numeric_cols[i],
                                'var2': numeric_cols[j],
                                'correlation': float(corr_val)
                            })
                
                results['correlation_analysis'] = {
                    'strong_correlations': strong_correlations
                }
            except Exception as e:
                print(f"相关性分析失败: {e}")
        
        # ========== 6. 关键发现 ==========
        findings = []
        
        # 数据量评估
        if len(df) > 10000:
            findings.append(f"数据集规模较大，共 {len(df):,} 条记录，具有较好的统计代表性")
        elif len(df) > 1000:
            findings.append(f"数据集共 {len(df):,} 条记录，数据量适中")
        else:
            findings.append(f"数据集共 {len(df):,} 条记录，建议谨慎推断结论")
        
        # 缺失值评估
        if results['data_quality']['columns_with_missing'] == 0:
            findings.append("数据完整性良好，无缺失值")
        else:
            findings.append(f"存在 {results['data_quality']['columns_with_missing']} 个字段有缺失值，建议关注数据质量")
        
        # 数值变量发现
        if numeric_cols and results['numeric_analysis'].get('statistics'):
            stats = results['numeric_analysis']['statistics']
            for col, stat in list(stats.items())[:3]:
                if stat['skewness'] > 1:
                    findings.append(f"【{col}】呈明显右偏分布（偏度={stat['skewness']:.2f}），大部分值集中在较低区间")
                elif stat['skewness'] < -1:
                    findings.append(f"【{col}】呈明显左偏分布（偏度={stat['skewness']:.2f}），大部分值集中在较高区间")
                
                cv = stat['std'] / stat['mean'] if stat['mean'] != 0 else 0
                if cv > 1:
                    findings.append(f"【{col}】变异系数较高（CV={cv:.2f}），数据离散程度大")
        
        # 分类变量发现
        if categorical_cols and results['categorical_analysis'].get('analysis'):
            cat = results['categorical_analysis']['analysis']
            for col, info in list(cat.items())[:2]:
                if info['unique_values'] == 1:
                    findings.append(f"【{col}】只有一个值，对分析无区分度")
                elif info['unique_values'] > 100:
                    findings.append(f"【{col}】有 {info['unique_values']} 个不同值，建议考虑分组或筛选")
        
        results['key_findings'] = findings
        
        return results
    
    def _generate_word_report(self, title: str, analysis: dict, output_path: Path, user_prompt: str = ""):
        """生成Word分析报告
        
        Args:
            title: 报告标题
            analysis: 分析结果字典
            output_path: 输出文件路径
            user_prompt: 用户分析提示词，用于定制报告内容和结构
        """
        
        doc = Document()
        
        # ========== 标题 ==========
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 如果有用户提示词，显示分析需求
        if user_prompt and user_prompt.strip():
            doc.add_paragraph("")
            doc.add_heading('分析需求', level=2)
            doc.add_paragraph(user_prompt.strip())
        
        doc.add_paragraph("")
        
        # ========== 一、数据概况 ==========
        doc.add_heading('一、数据概况', level=1)
        
        basic = analysis['basic_info']
        doc.add_paragraph(f"• 数据总量: {basic['total_rows']:,} 条记录")
        doc.add_paragraph(f"• 字段数量: {basic['total_columns']} 个")
        doc.add_paragraph(f"• 内存占用: {basic['memory_usage']}")
        doc.add_paragraph(f"• 重复记录: {basic['duplicate_rows']} 条")
        doc.add_paragraph("")
        
        # ========== 二、数据质量评估 ==========
        doc.add_heading('二、数据质量评估', level=1)
        
        quality = analysis['data_quality']
        doc.add_paragraph(f"• 数据完整率: {quality['completeness_rate']}")
        doc.add_paragraph(f"• 存在缺失的字段数: {quality['columns_with_missing']} 个")
        
        if quality['missing_details']:
            doc.add_paragraph("")
            doc.add_paragraph("缺失值详情:")
            for col, info in list(quality['missing_details'].items())[:5]:
                doc.add_paragraph(f"  - {col}: 缺失 {info['missing_count']} 条 ({info['missing_percentage']})")
        
        doc.add_paragraph("")
        
        # ========== 三、数值变量分析 ==========
        numeric_analysis = analysis.get('numeric_analysis', {})
        if numeric_analysis.get('numeric_columns') and numeric_analysis.get('statistics'):
            doc.add_heading('三、数值变量分析', level=1)
            
            doc.add_paragraph(f"数值型字段: {', '.join(numeric_analysis['numeric_columns'])}")
            doc.add_paragraph("")
            
            for col, stats in list(numeric_analysis['statistics'].items())[:5]:
                doc.add_heading(f'【{col}】', level=2)
                doc.add_paragraph(f"• 样本量: {stats['count']:,}")
                doc.add_paragraph(f"• 平均值: {stats['mean']:.2f}")
                doc.add_paragraph(f"• 中位数: {stats['median']:.2f}")
                doc.add_paragraph(f"• 标准差: {stats['std']:.2f}")
                doc.add_paragraph(f"• 范围: [{stats['min']:.2f}, {stats['max']:.2f}]")
                doc.add_paragraph(f"• 四分位数: Q1={stats['q1']:.2f}, Q3={stats['q3']:.2f}")
                
                if abs(stats['skewness']) > 0.5:
                    skew_desc = "右偏" if stats['skewness'] > 0 else "左偏"
                    doc.add_paragraph(f"• 分布形态: {skew_desc} (偏度={stats['skewness']:.2f})")
                
                # 插入图表
                chart_path = str(Path(output_path).parent.parent / 'charts' / f'numeric_{col}.png')
                if os.path.exists(chart_path):
                    try:
                        doc.add_picture(chart_path, width=Inches(5.5))
                    except:
                        pass
                
                doc.add_paragraph("")
        else:
            # 如果没有数值字段，给出说明
            doc.add_heading('三、数值变量分析', level=1)
            doc.add_paragraph("当前数据集不包含数值型字段，无法进行数值统计分析。")
            doc.add_paragraph("建议：如需进行数值分析，请上传包含数值字段的数据集。")
            doc.add_paragraph("")
        
        # ========== 四、分类变量分析 ==========
        if analysis.get('categorical_analysis'):
            doc.add_heading('四、分类变量分析', level=1)
            
            cat_analysis = analysis['categorical_analysis']
            doc.add_paragraph(f"分类型字段: {', '.join(cat_analysis['categorical_columns'])}")
            doc.add_paragraph("")
            
            if cat_analysis.get('analysis'):
                for col, info in list(cat_analysis['analysis'].items())[:4]:
                    doc.add_heading(f'【{col}】', level=2)
                    doc.add_paragraph(f"• 唯一值数量: {info['unique_values']}")
                    doc.add_paragraph("• 值分布 (Top 5):")
                    
                    for val, pct in list(info['top_percentage'].items())[:5]:
                        doc.add_paragraph(f"  - {val}: {pct}")
                    
                    # 插入图表
                    chart_path = str(Path(output_path).parent.parent / 'charts' / f'categorical_{col}.png')
                    if os.path.exists(chart_path):
                        try:
                            doc.add_picture(chart_path, width=Inches(5.5))
                        except:
                            pass
                    
                    doc.add_paragraph("")
        
        # ========== 五、相关性分析 ==========
        if analysis.get('correlation_analysis') and analysis['correlation_analysis'].get('strong_correlations'):
            doc.add_heading('五、相关性分析', level=1)
            
            corr = analysis['correlation_analysis']
            doc.add_paragraph("发现以下显著相关关系 (|r| > 0.5):")
            for item in corr['strong_correlations'][:5]:
                doc.add_paragraph(f"• {item['var1']} ↔ {item['var2']}: r = {item['correlation']:.3f}")
            
            # 插入热力图
            heatmap_path = str(Path(output_path).parent.parent / 'charts' / 'correlation_heatmap.png')
            if os.path.exists(heatmap_path):
                try:
                    doc.add_paragraph("")
                    doc.add_picture(heatmap_path, width=Inches(5.5))
                except:
                    pass
            
            doc.add_paragraph("")
        
        # ========== 六、LLM增强分析洞察（根据用户提示词定制）==========
        doc.add_heading('六、分析洞察', level=1)
        
        # 使用LLM生成深度分析（传入用户提示词）
        if llm_client.is_available():
            try:
                insights = self._generate_llm_insights(analysis, user_prompt)
                if insights:
                    doc.add_paragraph(insights)
                else:
                    # 使用关键发现
                    for finding in analysis.get('key_findings', []):
                        doc.add_paragraph(f"• {finding}")
            except Exception as e:
                # LLM失败时使用关键发现
                for finding in analysis.get('key_findings', []):
                    doc.add_paragraph(f"• {finding}")
        else:
            # 没有LLM时使用关键发现
            for finding in analysis.get('key_findings', []):
                doc.add_paragraph(f"• {finding}")
        
        doc.add_paragraph("")
        
        # ========== 七、分析建议（根据用户提示词定制）==========
        doc.add_heading('七、分析建议', level=1)
        
        # 使用LLM生成针对性建议（传入用户提示词）
        if llm_client.is_available():
            try:
                suggestions = self._generate_llm_suggestions(analysis, user_prompt)
                if suggestions:
                    doc.add_paragraph(suggestions)
                else:
                    self._add_default_suggestions(doc)
            except:
                self._add_default_suggestions(doc)
        else:
            self._add_default_suggestions(doc)
        
        doc.save(str(output_path))
    
    def _generate_llm_insights(self, analysis: dict, user_prompt: str = "") -> str:
        """使用LLM生成分析洞察
        
        Args:
            analysis: 分析结果字典
            user_prompt: 用户分析提示词，用于定制洞察方向
        """
        
        # 构建分析摘要
        summary = f"""
数据概况：
- 总记录数: {analysis['basic_info']['total_rows']:,}
- 字段数: {analysis['basic_info']['total_columns']}
- 数值字段数: {len(analysis.get('numeric_analysis', {}).get('numeric_columns', []))}
- 分类字段数: {len(analysis.get('categorical_analysis', {}).get('categorical_columns', []))}
- 数据完整率: {analysis['data_quality']['completeness_rate']}
"""
        
        # 添加分类变量关键信息
        if analysis.get('categorical_analysis', {}).get('analysis'):
            summary += "\n分类变量关键分布：\n"
            for col, info in list(analysis['categorical_analysis']['analysis'].items())[:3]:
                top_val = list(info['top_percentage'].items())[0] if info['top_percentage'] else ('N/A', '0%')
                summary += f"- {col}: 最主要类别是 '{top_val[0]}' 占比 {top_val[1]}\n"
        
        # 添加数值变量关键统计
        if analysis.get('numeric_analysis', {}).get('statistics'):
            summary += "\n数值变量关键统计：\n"
            for col, stats in list(analysis['numeric_analysis']['statistics'].items())[:3]:
                summary += f"- {col}: 平均值={stats['mean']:.2f}, 范围=[{stats['min']:.2f}, {stats['max']:.2f}]\n"
        
        # 根据用户提示词定制洞察方向
        prompt_instruction = "请基于以上数据分析结果，生成3-5条有价值的业务洞察。"
        
        if user_prompt and user_prompt.strip():
            # 从用户提示词中提取分析需求
            prompt_instruction = f"""
用户的分析需求如下：
{user_prompt.strip()}

请根据用户的分析需求，结合以上数据分析结果，生成针对性的业务洞察。
重点关注用户提到的分析维度和关注点。
"""
        
        prompt = f"""
{summary}

{prompt_instruction}

要求：
1. 每条洞察要具体、有数据支撑
2. 要有业务意义，不是简单的数据描述
3. 要指出潜在的问题或机会
4. 语言简洁专业
5. 如果用户有特定分析需求，优先响应这些需求

请直接输出洞察内容，每条一行，用数字编号。
"""
        
        try:
            response = llm_client.call(prompt)
            return response
        except:
            return ""
    
    def _generate_llm_suggestions(self, analysis: dict, user_prompt: str = "") -> str:
        """使用LLM生成针对性建议
        
        Args:
            analysis: 分析结果字典
            user_prompt: 用户分析提示词，用于定制建议方向
        """
        
        summary = f"""
数据集特征：
- {analysis['basic_info']['total_rows']:,} 条记录
- {analysis['basic_info']['total_columns']} 个字段
- 数据完整率: {analysis['data_quality']['completeness_rate']}
- 缺失字段数: {analysis['data_quality']['columns_with_missing']}
"""
        
        # 根据用户提示词定制建议方向
        prompt_instruction = "请针对这个数据集的特征，给出3-4条具体可行的分析或改进建议。"
        
        if user_prompt and user_prompt.strip():
            prompt_instruction = f"""
用户的分析需求如下：
{user_prompt.strip()}

请根据用户的分析需求，结合数据集特征，给出针对性的下一步分析建议或改进建议。
重点围绕用户关注的维度给出可执行的后续分析方案。
"""
        
        prompt = f"""
{summary}

{prompt_instruction}

要求：
1. 建议要针对性强，不是通用套话
2. 要考虑数据质量、分析价值、业务应用等维度
3. 语言简洁专业
4. 如果用户有特定分析需求，建议要与这些需求相关

请直接输出建议，每条一行，用数字编号。
"""
        
        try:
            response = llm_client.call(prompt)
            return response
        except:
            return ""
    
    def _add_default_suggestions(self, doc):
        """添加默认建议"""
        doc.add_paragraph("基于以上分析，建议:")
        doc.add_paragraph("1. 关注数据质量问题，对缺失值进行合理处理")
        doc.add_paragraph("2. 对关键指标进行深入的趋势和对比分析")
        doc.add_paragraph("3. 探索变量之间的关系和业务逻辑")
        doc.add_paragraph("4. 结合业务场景解读数据特征")