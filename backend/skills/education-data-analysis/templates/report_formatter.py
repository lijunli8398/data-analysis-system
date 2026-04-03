#!/usr/bin/env python3
"""
Word报告格式化模块

提供专业的Word文档格式化功能，包括：
- 标准化的标题样式
- 规范的表格格式
- 图表标题和编号
- 段落间距和对齐
- 页眉页脚设置
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


class ReportFormatter:
    """Word报告格式化器（支持自动编号）"""
    
    def __init__(self, doc=None):
        """初始化格式化器
        
        Args:
            doc: Document对象，如果为None则创建新文档
        """
        self.doc = doc if doc else Document()
        self._figure_counter = 0  # 图表计数器
        self._table_counter = 0   # 表格计数器
        self._setup_styles()
        self._setup_page_settings()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置正文样式
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(6)
        
        # 设置标题1样式
        style = self.doc.styles['Heading 1']
        style.font.name = '黑体'
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 设置标题2样式
        style = self.doc.styles['Heading 2']
        style.font.name = '黑体'
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 设置标题3样式
        style = self.doc.styles['Heading 3']
        style.font.name = '黑体'
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(2)
    
    def _setup_page_settings(self):
        """设置页面格式"""
        section = self.doc.sections[0]
        # 页边距
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    def add_cover_page(self, title, subtitle=None, date=None, organization=None):
        """添加封面页
        
        Args:
            title: 报告标题
            subtitle: 副标题（可选）
            date: 日期（可选，默认当前日期）
            organization: 组织名称（可选）
        """
        # 添加空行调整位置
        for _ in range(3):
            self.doc.add_paragraph()
        
        # 主标题
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.name = '黑体'
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 副标题
        if subtitle:
            self.doc.add_paragraph()
            sub_para = self.doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.name = '宋体'
            sub_run.font.size = Pt(16)
            sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加空行
        for _ in range(5):
            self.doc.add_paragraph()
        
        # 日期
        if date is None:
            date = datetime.datetime.now().strftime('%Y年%m月%d日')
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(date)
        date_run.font.name = '宋体'
        date_run.font.size = Pt(14)
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 组织名称
        if organization:
            self.doc.add_paragraph()
            org_para = self.doc.add_paragraph()
            org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            org_run = org_para.add_run(organization)
            org_run.font.name = '宋体'
            org_run.font.size = Pt(14)
            org_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 分页
        self.doc.add_page_break()
    
    def add_section_heading(self, text, level=1):
        """添加章节标题
        
        Args:
            text: 标题文本
            level: 标题级别（1-3）
        """
        heading = self.doc.add_heading(text, level=level)
        return heading
    
    def add_paragraph_with_style(self, text, bold_prefix=None, indent=False):
        """添加格式化段落
        
        Args:
            text: 段落内容
            bold_prefix: 加粗的前缀文字（可选）
            indent: 是否首行缩进
        """
        para = self.doc.add_paragraph()
        
        if indent:
            para.paragraph_format.first_line_indent = Cm(0.74)
        
        if bold_prefix:
            bold_run = para.add_run(bold_prefix)
            bold_run.font.bold = True
            bold_run.font.name = '宋体'
            bold_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        content_run = para.add_run(text)
        content_run.font.name = '宋体'
        content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        return para
    
    def add_bullet_point(self, text, bold_prefix=None):
        """添加要点列表项
        
        Args:
            text: 要点内容
            bold_prefix: 加粗的前缀（可选）
        """
        para = self.doc.add_paragraph(style='List Bullet')
        
        if bold_prefix:
            bold_run = para.add_run(bold_prefix)
            bold_run.font.bold = True
            bold_run.font.name = '宋体'
            bold_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        content_run = para.add_run(text)
        content_run.font.name = '宋体'
        content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        return para
    
    def add_figure(self, image_path, title, width=Inches(6)):
        """添加图表（自动编号）
        
        Args:
            image_path: 图片路径
            title: 图表标题
            width: 图片宽度
        """
        self._figure_counter += 1
        figure_number = f"图{self._figure_counter}"
        
        # 添加图片
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if str(image_path).endswith('.png') or str(image_path).endswith('.jpg'):
            para.add_run().add_picture(str(image_path), width=width)
        
        # 添加图表标题
        caption_para = self.doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        caption_text = f"{figure_number}  {title}"
        
        caption_run = caption_para.add_run(caption_text)
        caption_run.font.name = '宋体'
        caption_run.font.size = Pt(10)
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 图表标题后添加空行
        self.doc.add_paragraph()
        
        return para
    
    def add_table_with_header(self, headers, data, title=None):
        """添加格式化表格（自动编号）
        
        Args:
            headers: 表头列表
            data: 数据行列表（每行是一个列表）
            title: 表格标题（可选）
        """
        self._table_counter += 1
        table_number = f"表{self._table_counter}"
        
        # 表格标题
        if title:
            caption_para = self.doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            caption_text = f"{table_number}  {title}"
            
            caption_run = caption_para.add_run(caption_text)
            caption_run.font.name = '黑体'
            caption_run.font.size = Pt(10)
            caption_run.font.bold = True
            caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 创建表格
        table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            
            # 表头加粗居中
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '宋体'
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 表头背景色（浅灰）
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E7E6E6')
            cell._tc.get_or_add_tcPr().append(shading)
        
        # 填充数据
        for row_idx, row_data in enumerate(data):
            row = table.rows[row_idx + 1]
            for col_idx, cell_data in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = str(cell_data)
                
                for paragraph in cell.paragraphs:
                    # 数值居中，文本左对齐
                    try:
                        float(cell_data)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 表格后添加空行
        self.doc.add_paragraph()
        
        return table
    
    def add_key_findings(self, findings, add_heading=False):
        """添加关键发现模块
        
        Args:
            findings: 发现列表，每个元素是 (标题, 内容) 元组
            add_heading: 是否添加标题（默认False，由外层调用处添加）
        """
        if add_heading:
            heading = self.add_section_heading('关键发现', level=2)
        
        for title, content in findings:
            para = self.doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(0.74)
            
            title_run = para.add_run(f"{title}：")
            title_run.font.bold = True
            title_run.font.name = '宋体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            content_run = para.add_run(content)
            content_run.font.name = '宋体'
            content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        self.doc.add_paragraph()
    
    def add_analysis_summary(self, summary_points, add_heading=False):
        """添加分析总结
        
        Args:
            summary_points: 总结要点列表
            add_heading: 是否添加标题（默认False，由外层调用处添加）
        """
        if add_heading:
            heading = self.add_section_heading('分析总结', level=2)
        
        for point in summary_points:
            para = self.doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(0.74)
            
            run = para.add_run(point)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        self.doc.add_paragraph()
    
    def add_recommendations(self, recommendations, add_heading=False):
        """添加建议模块
        
        Args:
            recommendations: 建议列表，每个元素包含类别和建议内容
            add_heading: 是否添加标题（默认False，由外层调用处添加）
        """
        if add_heading:
            heading = self.add_section_heading('改进建议', level=2)
        
        for i, (category, items) in enumerate(recommendations, 1):
            # 类别标题（使用（1）（2）（3）格式）
            cat_para = self.doc.add_paragraph()
            cat_run = cat_para.add_run(f"（{i}）{category}")
            cat_run.font.bold = True
            cat_run.font.name = '宋体'
            cat_run.font.size = Pt(12)
            cat_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 建议内容
            for item in items:
                self.add_bullet_point(item)
            
            self.doc.add_paragraph()
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def add_toc_placeholder(self):
        """添加目录占位（Word需手动更新目录）"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run('目  录')
        run.font.name = '黑体'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        self.doc.add_paragraph()
        
        # 添加目录字段（需要用户在Word中更新）
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        # 目录说明
        note_para = self.doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note_para.add_run('（请在Word中右键点击更新目录）')
        note_run.font.name = '宋体'
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(128, 128, 128)
        note_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        self.add_page_break()
    
    def get_document(self):
        """获取Document对象"""
        return self.doc
    
    def save(self, path):
        """保存文档
        
        Args:
            path: 保存路径
        """
        self.doc.save(str(path))


def create_formatted_report(title='教育学情数据分析报告'):
    """创建格式化的报告文档
    
    Args:
        title: 报告标题
    
    Returns:
        ReportFormatter对象
    """
    formatter = ReportFormatter()
    
    # 添加封面
    formatter.add_cover_page(
        title=title,
        subtitle='数据分析与干预建议',
        organization='教育数据分析中心'
    )
    
    # 不再添加目录页（直接进入正文）
    
    return formatter