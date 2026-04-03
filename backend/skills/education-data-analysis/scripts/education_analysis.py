#!/usr/bin/env python3
"""
教育学情数据分析脚本（动态版本 - 不硬编码数据结构）

Usage:
    python education_analysis.py --data <data_file> [--output <output_dir>] [--verbose]
    
Example:
    python education_analysis.py --data data/学生指标得分清单.xlsx
    python education_analysis.py --data data/学生指标得分清单.xlsx --output custom_output --verbose
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import sys
import json
import datetime
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

# 导入报告格式化模块
templates_path = Path(__file__).parent.parent / 'templates'
if str(templates_path) not in sys.path:
    sys.path.insert(0, str(templates_path))

try:
    from report_formatter import ReportFormatter, create_formatted_report
    HAS_FORMATTER = True
except ImportError as e:
    HAS_FORMATTER = False
    print(f"警告：报告格式化模块导入失败，将使用基本格式: {e}")

# 设置中文字体（使用系统已有的中文字体）
# 注意：Noto Sans CJK 在 .ttc 文件中可能被识别为 JP 变体，使用 WenQuanYi 更可靠
plt.rcParams['font.sans-serif'] = ['WenQuanYi Micro Hei', 'Noto Sans CJK SC', 'SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['figure.dpi'] = 300


class EducationDataAnalyzer:
    """教育学情数据分析器（动态版本 - 自动推断数据结构）"""
    
    def __init__(self, data_path, output_dir='output', env_prefix='成长环境_', dev_prefix='学生发展_'):
        self.data_path = Path(data_path)
        self.output_dir = Path(output_dir)
        self.df = None
        self.df_clean = None
        
        # 指标前缀配置（可通过参数修改）
        self.env_prefix = env_prefix
        self.dev_prefix = dev_prefix
        
        # 指标名称映射（动态生成）
        self.indicator_names = {}
        
        # 指标列表（动态推断）
        self.env_indicators = []
        self.dev_indicators = []
        self.all_indicators = []
        
        # 基础信息列（动态推断）
        self.base_columns = []
        
        # 分类值（动态推断）
        self.gender_values = []
        self.scale_values = []
        self.nature_values = []
        
        # 风险阈值
        self.risk_threshold = 60
        
        # 创建输出目录
        self._create_output_dirs()
        
    def _create_output_dirs(self):
        """创建输出目录结构"""
        dirs = ['data', 'report', 'report_charts']
        for dir_name in dirs:
            (self.output_dir / dir_name).mkdir(parents=True, exist_ok=True)
    
    def _detect_header_row(self, df_raw):
        """自动检测表头行数和结构
        
        支持三种表头格式：
        1. 单行表头（包含成长环境_、学生发展_前缀）
        2. 三级表头（成长环境/学生发展在第一行，具体指标在第三行）
        """
        # 检查是否有三级表头结构
        # 第一行包含"成长环境"或"学生发展"（不带下划线）
        row0 = df_raw.iloc[0].astype(str).tolist()
        has_multi_level = any('成长环境' in str(cell) or '学生发展' in str(cell) for cell in row0)
        
        if has_multi_level:
            # 检查第三行是否是具体指标名
            row2 = df_raw.iloc[2].astype(str).tolist() if len(df_raw) > 2 else []
            # 第三行应该有具体指标（如亲子关系、师生关系等）
            if row2 and any(str(cell) not in ['办学规模', '办学性质', '学校代码', '年级', '班级', '考号', '姓名', '性别', '成长环境', '学生发展', 'nan', ''] for cell in row2[8:]):
                return 2, 'multi_level'  # 三级表头，数据从第3行开始
        
        # 单行表头：查找包含指标前缀（带下划线）的行
        for i in range(min(10, len(df_raw))):
            row = df_raw.iloc[i].astype(str).tolist()
            if any(self.env_prefix in str(cell) or self.dev_prefix in str(cell) for cell in row):
                return i, 'single'
        
        # 默认跳过3行（兼容旧数据格式）
        return 3, 'single'
    
    def _detect_columns(self, df_raw, header_row, header_type='single'):
        """从表头自动推断列名
        
        对于三级表头，需要合并生成正确的列名：
        成长环境 + 亲子关系 = 成长环境_亲子关系
        """
        if header_type == 'multi_level':
            # 三级表头合并处理
            columns = []
            for i in range(len(df_raw.columns)):
                level0 = str(df_raw.iloc[0, i]).strip() if pd.notna(df_raw.iloc[0, i]) else ''
                level1 = str(df_raw.iloc[1, i]).strip() if pd.notna(df_raw.iloc[1, i]) else ''
                level2 = str(df_raw.iloc[2, i]).strip() if pd.notna(df_raw.iloc[2, i]) else ''
                
                # 基础信息列（前8列）
                if i < 8:
                    col = level2 if level2 and level2 not in ['办学规模', '办学性质', '学校代码', '年级', '班级', '考号', '姓名', '性别'] else \
                          level1 if level1 and level1 not in ['办学规模', '办学性质', '学校代码', '年级', '班级', '考号', '姓名', '性别'] else level0
                else:
                    # 指标列：合并第一行和第三行
                    if level0 in ['成长环境', '学生发展']:
                        col = f"{level0}_{level2}"
                    elif level0.startswith('成长环境') or level0.startswith('学生发展'):
                        # 已经是合并格式
                        col = level0
                    else:
                        col = level2 if level2 else level1 if level1 else level0
                
                columns.append(col if col else f'未命名_{i}')
        else:
            # 单行表头
            columns = df_raw.iloc[header_row].tolist()
            columns = [str(col).strip() if pd.notna(col) else f'未命名_{i}' for i, col in enumerate(columns)]
        
        return columns
    
    def _detect_indicators(self, columns):
        """从列名中自动推断指标列表"""
        # 成长环境指标
        self.env_indicators = [col for col in columns if self.env_prefix in col]
        
        # 学生发展指标  
        self.dev_indicators = [col for col in columns if self.dev_prefix in col]
        
        # 所有指标
        self.all_indicators = self.env_indicators + self.dev_indicators
        
        # 基础信息列（排除指标列）
        self.base_columns = [col for col in columns 
                            if self.env_prefix not in col and self.dev_prefix not in col 
                            and col not in ['成长环境_均分', '学生发展_均分']]
        
        # 生成指标名称映射（去除前缀作为简称）
        for col in self.all_indicators:
            if self.env_prefix in col:
                short_name = col.replace(self.env_prefix, '')
            elif self.dev_prefix in col:
                short_name = col.replace(self.dev_prefix, '')
            else:
                short_name = col
            self.indicator_names[col] = short_name
        
        return self.env_indicators, self.dev_indicators
    
    def _detect_categories(self, df):
        """从数据中自动推断分类值"""
        # 性别分类
        if '性别' in df.columns:
            self.gender_values = df['性别'].dropna().unique().tolist()
            # 过滤掉脏数据
            self.gender_values = [v for v in self.gender_values 
                                 if str(v) not in ['办学规模', '办学性质', '性别', 'nan', '']]
        
        # 办学规模分类
        if '办学规模' in df.columns:
            self.scale_values = df['办学规模'].dropna().unique().tolist()
            self.scale_values = [v for v in self.scale_values 
                                if str(v) not in ['办学规模', '办学性质', '性别', 'nan', '']]
        
        # 办学性质分类
        if '办学性质' in df.columns:
            self.nature_values = df['办学性质'].dropna().unique().tolist()
            self.nature_values = [v for v in self.nature_values 
                                 if str(v) not in ['办学规模', '办学性质', '性别', 'nan', '']]
        
        return self.gender_values, self.scale_values, self.nature_values
    
    def load_data(self, verbose=False):
        """加载并清洗数据"""
        if verbose:
            print(f"正在加载数据: {self.data_path}")
        
        # 读取原始数据
        df_raw = pd.read_excel(self.data_path, header=None)
        
        if verbose:
            print(f"原始数据维度: {df_raw.shape}")
        
        # 自动检测表头行和类型
        header_row, header_type = self._detect_header_row(df_raw)
        if verbose:
            print(f"检测到表头行: 第{header_row}行, 类型: {header_type}")
        
        # 自动推断列名
        columns = self._detect_columns(df_raw, header_row, header_type)
        
        if verbose:
            print(f"检测到列名: {columns}")
        
        # 自动推断指标列表
        self._detect_indicators(columns)
        
        if verbose:
            print(f"检测到成长环境指标: {self.env_indicators}")
            print(f"检测到学生发展指标: {self.dev_indicators}")
        
        # 截取数据（跳过表头行）
        df = df_raw.iloc[header_row + 1:].copy()
        df.columns = columns
        df = df.reset_index(drop=True)
        
        # 转换数值列
        for col in self.all_indicators:
            df[col] = pd.to_numeric(df[col], errors='coerce')
        
        # 计算综合得分（动态）
        if len(self.env_indicators) > 0:
            df['成长环境_均分'] = df[self.env_indicators].mean(axis=1)
        else:
            df['成长环境_均分'] = np.nan
            
        if len(self.dev_indicators) > 0:
            df['学生发展_均分'] = df[self.dev_indicators].mean(axis=1)
        else:
            df['学生发展_均分'] = np.nan
        
        # 清理脏数据（动态检测需要排除的值）
        dirty_values = ['办学规模', '办学性质', '性别', 'nan', '', 'NaN']
        
        for col in df.columns:
            if col in ['办学规模', '办学性质', '性别']:
                df = df[~df[col].isin(dirty_values)]
        
        # 删除指标列含缺失值的行
        valid_indicator_cols = [col for col in self.all_indicators if col in df.columns]
        df_clean = df.dropna(subset=valid_indicator_cols + ['成长环境_均分', '学生发展_均分'])
        
        # 自动推断分类值
        self._detect_categories(df_clean)
        
        if verbose:
            print(f"检测到性别分类: {self.gender_values}")
            print(f"检测到办学规模分类: {self.scale_values}")
            print(f"检测到办学性质分类: {self.nature_values}")
        
        self.df = df
        self.df_clean = df_clean
        
        if verbose:
            print(f"清洗后数据量: {len(df_clean)}条记录")
            if '学校代码' in df_clean.columns:
                print(f"学校数量: {df_clean['学校代码'].nunique()}所")
        
        # 保存清洗后的数据和配置信息
        df_clean.to_csv(self.output_dir / 'data' / 'analysis_data_clean.csv', 
                        index=False, encoding='utf-8-sig')
        
        # 保存数据配置信息
        config = {
            'env_prefix': self.env_prefix,
            'dev_prefix': self.dev_prefix,
            'env_indicators': self.env_indicators,
            'dev_indicators': self.dev_indicators,
            'indicator_names': self.indicator_names,
            'gender_values': self.gender_values,
            'scale_values': self.scale_values,
            'nature_values': self.nature_values,
            'header_row': header_row
        }
        
        with open(self.output_dir / 'data' / 'data_config.json', 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        
        return df_clean
    
    def analyze_overall(self, verbose=False):
        """整体学情分析"""
        if verbose:
            print("\n="*60)
            print("二、整体学情画像分析")
            print("="*60)
        
        df = self.df_clean
        
        # 1. 数据概况
        total_students = len(df)
        total_schools = df['学校代码'].nunique() if '学校代码' in df.columns else 0
        
        stats_summary = {
            '总学生数': total_students,
            '学校数量': total_schools
        }
        
        # 性别分布（动态）
        if '性别' in df.columns and len(self.gender_values) > 0:
            gender_dist = df['性别'].value_counts()
            for gender in self.gender_values:
                stats_summary[f'{gender}生数'] = gender_dist.get(gender, 0)
            
            if verbose:
                print(f"\n性别分布：")
                for gender in self.gender_values:
                    count = gender_dist.get(gender, 0)
                    print(f"  {gender}: {count}人 ({count/total_students*100:.2f}%)")
        
        # 办学性质分布（动态）
        if '办学性质' in df.columns and len(self.nature_values) > 0:
            nature_dist = df['办学性质'].value_counts()
            for nature in self.nature_values:
                stats_summary[f'{nature}学校学生数'] = nature_dist.get(nature, 0)
            
            if verbose:
                print(f"\n办学性质分布：")
                for nature in self.nature_values:
                    count = nature_dist.get(nature, 0)
                    print(f"  {nature}: {count}人 ({count/total_students*100:.2f}%)")
        
        if verbose:
            print(f"\n总样本量：{total_students}名学生")
            if total_schools > 0:
                print(f"学校数量：{total_schools}所学校")
        
        # 保存基本统计
        with open(self.output_dir / 'data' / 'stats_summary.txt', 'w', encoding='utf-8') as f:
            f.write("教育学情数据基本统计\n")
            f.write("="*60 + "\n\n")
            f.write(f"总样本量：{total_students}名学生\n")
            if total_schools > 0:
                f.write(f"学校数量：{total_schools}所学校\n\n")
            
            if '性别' in df.columns:
                f.write("性别分布：\n")
                for gender in self.gender_values:
                    count = df['性别'].value_counts().get(gender, 0)
                    f.write(f"  {gender}：{count}人 ({count/total_students*100:.2f}%)\n")
                f.write("\n")
            
            if '办学性质' in df.columns:
                f.write("办学性质分布：\n")
                for nature in self.nature_values:
                    count = df['办学性质'].value_counts().get(nature, 0)
                    f.write(f"  {nature}：{count}人 ({count/total_students*100:.2f}%)\n")
                f.write("\n")
            
            if '办学规模' in df.columns:
                f.write("办学规模分布：\n")
                for scale in self.scale_values:
                    count = df['办学规模'].value_counts().get(scale, 0)
                    f.write(f"  {scale}：{count}人 ({count/total_students*100:.2f}%)\n")
        
        # 2. 各指标统计（动态）
        if len(self.all_indicators) == 0:
            if verbose:
                print("警告：未检测到指标列，无法进行指标分析")
            return stats_summary, pd.DataFrame()
        
        avg_scores = df[self.all_indicators].mean()
        risk_rates = {}
        
        for col in self.all_indicators:
            if col in df.columns:
                risk_rate = (df[col] < self.risk_threshold).sum() / len(df) * 100
                risk_rates[col] = risk_rate
        
        # 保存指标统计
        indicator_stats = pd.DataFrame({
            '指标': [self.indicator_names.get(col, col) for col in self.all_indicators],
            '完整列名': self.all_indicators,
            '平均分': [avg_scores.get(col, np.nan) for col in self.all_indicators],
            '风险暴露率(%)': [risk_rates.get(col, np.nan) for col in self.all_indicators]
        })
        
        indicator_stats.to_csv(self.output_dir / 'data' / 'indicator_stats.csv', 
                                index=False, encoding='utf-8-sig')
        
        if verbose:
            print("\n各指标平均分：")
            for col in self.all_indicators:
                if col in avg_scores:
                    print(f"  {self.indicator_names.get(col, col)}: {avg_scores[col]:.2f}")
            
            print("\n各指标风险暴露率（低于60分）：")
            for col in self.all_indicators:
                if col in risk_rates:
                    print(f"  {self.indicator_names.get(col, col)}: {risk_rates[col]:.2f}%")
        
        return stats_summary, indicator_stats
    
    def analyze_group_differences(self, verbose=False):
        """群体差异分析（动态版本）"""
        if verbose:
            print("\n="*60)
            print("三、群体差异分析")
            print("="*60)
        
        df = self.df_clean
        
        # 1. 性别差异（动态）
        gender_diff = pd.DataFrame()
        
        if '性别' in df.columns and len(self.gender_values) > 0:
            for gender in self.gender_values:
                gender_data = df[df['性别'] == gender]
                if len(gender_data) > 0:
                    row = {
                        '性别': gender,
                        '人数': len(gender_data),
                        '成长环境均分': gender_data['成长环境_均分'].mean(),
                        '学生发展均分': gender_data['学生发展_均分'].mean()
                    }
                    
                    # 添加各指标（动态）
                    for col in self.all_indicators:
                        if col in gender_data.columns:
                            row[self.indicator_names.get(col, col)] = gender_data[col].mean()
                    
                    gender_diff = pd.concat([gender_diff, pd.DataFrame([row])], ignore_index=True)
            
            gender_diff.to_csv(self.output_dir / 'data' / 'gender_analysis.csv', 
                               index=False, encoding='utf-8-sig')
            
            if verbose:
                print("\n性别差异分析：")
                print(gender_diff)
        else:
            if verbose:
                print("\n未检测到性别数据，跳过性别差异分析")
        
        # 2. 办学规模差异（动态）
        scale_diff = pd.DataFrame()
        
        if '办学规模' in df.columns and len(self.scale_values) > 0:
            for scale in self.scale_values:
                scale_data = df[df['办学规模'] == scale]
                if len(scale_data) > 0:
                    row = {
                        '办学规模': scale,
                        '人数': len(scale_data),
                        '成长环境均分': scale_data['成长环境_均分'].mean(),
                        '学生发展均分': scale_data['学生发展_均分'].mean()
                    }
                    scale_diff = pd.concat([scale_diff, pd.DataFrame([row])], ignore_index=True)
            
            scale_diff.to_csv(self.output_dir / 'data' / 'scale_analysis.csv', 
                              index=False, encoding='utf-8-sig')
            
            if verbose:
                print("\n办学规模差异分析：")
                print(scale_diff)
        else:
            if verbose:
                print("\n未检测到办学规模数据，跳过办学规模差异分析")
        
        # 3. 办学性质差异（动态）
        nature_diff = pd.DataFrame()
        
        if '办学性质' in df.columns and len(self.nature_values) > 0:
            for nature in self.nature_values:
                nature_data = df[df['办学性质'] == nature]
                if len(nature_data) > 0:
                    row = {
                        '办学性质': nature,
                        '人数': len(nature_data),
                        '成长环境均分': nature_data['成长环境_均分'].mean(),
                        '学生发展均分': nature_data['学生发展_均分'].mean()
                    }
                    nature_diff = pd.concat([nature_diff, pd.DataFrame([row])], ignore_index=True)
            
            nature_diff.to_csv(self.output_dir / 'data' / 'nature_analysis.csv', 
                               index=False, encoding='utf-8-sig')
            
            if verbose:
                print("\n办学性质差异分析：")
                print(nature_diff)
        else:
            if verbose:
                print("\n未检测到办学性质数据，跳过办学性质差异分析")
        
        return gender_diff, scale_diff, nature_diff
    
    def analyze_school_level(self, verbose=False):
        """学校层面分析"""
        if verbose:
            print("\n="*60)
            print("四、学校层面分析")
            print("="*60)
        
        df = self.df_clean
        
        if '学校代码' not in df.columns:
            if verbose:
                print("未检测到学校代码列，跳过学校层面分析")
            return pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
        
        # 计算各学校平均分
        agg_dict = {
            '成长环境_均分': 'mean',
            '学生发展_均分': 'mean'
        }
        if '考号' in df.columns:
            agg_dict['考号'] = 'count'
        elif '姓名' in df.columns:
            agg_dict['姓名'] = 'count'
        
        school_stats = df.groupby('学校代码').agg(agg_dict).round(2)
        
        # 重命名人数列
        if '考号' in school_stats.columns:
            school_stats = school_stats.rename(columns={'考号': '学生人数'})
        elif '姓名' in school_stats.columns:
            school_stats = school_stats.rename(columns={'姓名': '学生人数'})
        
        school_stats = school_stats.sort_values('学生发展_均分', ascending=False)
        
        # 前5名学校
        top5 = school_stats.head(5)
        
        # 后5名学校
        bottom5 = school_stats.tail(5)
        
        if verbose:
            print("\n发展均分较高的学校（前5）：")
            print(top5)
            
            print("\n发展均分较低的学校（后5）：")
            print(bottom5)
        
        # 保存学校数据
        school_stats.to_csv(self.output_dir / 'data' / 'school_analysis.csv', 
                            encoding='utf-8-sig')
        
        # 学校定位数据
        school_position = school_stats.reset_index()
        school_position.to_csv(self.output_dir / 'data' / 'school_position.csv', 
                                index=False, encoding='utf-8-sig')
        
        return school_stats, top5, bottom5
    
    def analyze_quadrant(self, verbose=False):
        """四象限分层分析"""
        if verbose:
            print("\n="*60)
            print("五、学生分层分析")
            print("="*60)
        
        df = self.df_clean
        
        # 检查是否有必要的均分列
        if '成长环境_均分' not in df.columns or '学生发展_均分' not in df.columns:
            if verbose:
                print("缺少均分列，无法进行四象限分析")
            return pd.Series(), pd.DataFrame(), pd.DataFrame()
        
        # 计算中位数（动态）
        env_median = df['成长环境_均分'].median()
        dev_median = df['学生发展_均分'].median()
        
        if verbose:
            print(f"\n成长环境中位数：{env_median:.2f}")
            print(f"学生发展中位数：{dev_median:.2f}")
        
        # 分配象限
        def classify_quadrant(row):
            env = row['成长环境_均分']
            dev = row['学生发展_均分']
            
            if pd.isna(env) or pd.isna(dev):
                return '未知'
            
            if env >= env_median and dev >= dev_median:
                return '第一象限-优势发展型'
            elif env < env_median and dev >= dev_median:
                return '第二象限-逆境成长型'
            elif env >= env_median and dev < dev_median:
                return '第三象限-潜力待发型'
            else:
                return '第四象限-重点关注型'
        
        df['象限'] = df.apply(classify_quadrant, axis=1)
        
        # 统计各象限人数
        quadrant_stats = df['象限'].value_counts()
        
        if verbose:
            print("\n四象限学生分布：")
            for quadrant, count in quadrant_stats.items():
                if quadrant != '未知':
                    print(f"  {quadrant}: {count}人 ({count/len(df)*100:.2f}%)")
        
        # 各象限详细数据
        quadrant_detailed = df[df['象限'] != '未知'].groupby('象限').agg({
            '成长环境_均分': 'mean',
            '学生发展_均分': 'mean'
        }).round(2)
        
        # 添加人数
        quadrant_detailed['人数'] = df[df['象限'] != '未知'].groupby('象限').size()
        
        quadrant_detailed.to_csv(self.output_dir / 'data' / 'quadrant_analysis.csv', 
                                  encoding='utf-8-sig')
        
        # 生成四象限干预建议
        intervention_suggestions = pd.DataFrame([
            {
                '象限': '第一象限-优势发展型',
                '人数': quadrant_stats.get('第一象限-优势发展型', 0),
                '特征': '成长环境良好，学生发展表现优秀',
                '干预建议': '继续保持优势，提供拓展性学习资源，培养领导力和创新能力，鼓励参与社会实践和竞赛活动'
            },
            {
                '象限': '第二象限-逆境成长型',
                '人数': quadrant_stats.get('第二象限-逆境成长型', 0),
                '特征': '成长环境欠佳，但学生发展表现优秀',
                '干预建议': '关注学生心理韧性，提供情感支持，改善成长环境，建立师生一对一帮扶机制，预防环境因素导致的潜在风险'
            },
            {
                '象限': '第三象限-潜力待发型',
                '人数': quadrant_stats.get('第三象限-潜力待发型', 0),
                '特征': '成长环境良好，但学生发展表现欠佳',
                '干预建议': '深入了解学生个体差异，提供个性化学习指导，激发学习兴趣，建立激励机制，探索适合的教学方法'
            },
            {
                '象限': '第四象限-重点关注型',
                '人数': quadrant_stats.get('第四象限-重点关注型', 0),
                '特征': '成长环境欠佳，学生发展表现欠佳',
                '干预建议': '优先干预，建立多维度帮扶体系，改善家庭和校园环境，提供心理辅导，制定个性化发展计划，定期跟踪评估'
            }
        ])
        
        intervention_suggestions.to_csv(self.output_dir / 'data' / 'intervention_suggestions.csv', 
                                        index=False, encoding='utf-8-sig')
        
        return quadrant_stats, quadrant_detailed, intervention_suggestions
    
    def generate_charts(self, verbose=False):
        """生成所有图表（动态版本）"""
        if verbose:
            print("\n="*60)
            print("生成可视化图表")
            print("="*60)
        
        charts_dir = self.output_dir / 'report_charts'
        
        # 读取数据
        indicator_stats_path = self.output_dir / 'data' / 'indicator_stats.csv'
        if not indicator_stats_path.exists():
            if verbose:
                print("警告：缺少指标统计数据，无法生成图表")
            return
        
        indicator_stats = pd.read_csv(indicator_stats_path)
        
        # 图1: 各指标平均分对比（动态）
        if len(indicator_stats) > 0:
            fig, ax = plt.subplots(figsize=(max(10, len(indicator_stats) * 1.2), 6))
            # 颜色区分：成长环境用蓝色，学生发展用橙色
            colors = ['#4472C4' if col in self.env_indicators else '#ED7D31' 
                      for col in indicator_stats['完整列名']]
            bars = ax.bar(indicator_stats['指标'], indicator_stats['平均分'], color=colors)
            ax.set_ylabel('平均分', fontsize=12)
            ax.set_title('各指标平均分对比', fontsize=14, fontweight='bold')
            ax.set_ylim(0, 100)
            ax.axhline(y=60, color='red', linestyle='--', linewidth=1, label='及格线(60分)')
            ax.legend()
            
            for bar, value in zip(bars, indicator_stats['平均分']):
                if pd.notna(value):
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1, 
                            f'{value:.1f}', ha='center', va='bottom', fontsize=9)
            
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            plt.savefig(charts_dir / '各指标平均分对比.png', bbox_inches='tight')
            plt.close()
            
            if verbose:
                print("图表：各指标平均分对比 已保存")
        
        # 图2: 各指标风险暴露率（动态）
        if len(indicator_stats) > 0:
            fig, ax = plt.subplots(figsize=(max(10, len(indicator_stats) * 1.2), 6))
            risk_values = indicator_stats['风险暴露率(%)'].fillna(0)
            colors = ['#FF6B6B' if rate > 40 else '#FFA07A' if rate > 20 else '#90EE90' 
                      for rate in risk_values]
            bars = ax.bar(indicator_stats['指标'], risk_values, color=colors)
            ax.set_ylabel('风险暴露率 (%)', fontsize=12)
            ax.set_title('各指标风险暴露率（低于60分）', fontsize=14, fontweight='bold')
            ax.set_ylim(0, max(60, max(risk_values) + 10))
            
            for bar, value in zip(bars, risk_values):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                        f'{value:.1f}%', ha='center', va='bottom', fontsize=9)
            
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            plt.savefig(charts_dir / '各指标风险暴露率.png', bbox_inches='tight')
            plt.close()
            
            if verbose:
                print("图表：各指标风险暴露率 已保存")
        
        # 图3: 性别差异对比（动态）
        gender_analysis_path = self.output_dir / 'data' / 'gender_analysis.csv'
        if gender_analysis_path.exists() and len(self.gender_values) >= 2:
            gender_analysis = pd.read_csv(gender_analysis_path)
            
            if len(gender_analysis) >= 2:
                fig, axes = plt.subplots(1, 2, figsize=(14, 5))
                
                # 成长环境指标（动态）
                env_indicators_short = [self.indicator_names.get(col, col) 
                                        for col in self.env_indicators]
                x = np.arange(len(env_indicators_short))
                width = 0.35
                
                # 获取各性别数据
                for i, gender in enumerate(self.gender_values[:2]):
                    gender_data = gender_analysis[gender_analysis['性别'] == gender]
                    if len(gender_data) > 0:
                        values = [gender_data.get(short, 0).values[0] if short in gender_data.columns else 0
                                  for short in env_indicators_short]
                        if i == 0:
                            bars1 = axes[0].bar(x - width/2, values, width, 
                                                label=self.gender_values[0], color='#4472C4')
                        else:
                            bars2 = axes[0].bar(x + width/2, values, width, 
                                                label=self.gender_values[1], color='#ED7D31')
                
                axes[0].set_ylabel('平均分', fontsize=11)
                axes[0].set_title('成长环境维度性别差异', fontsize=12, fontweight='bold')
                axes[0].set_xticks(x)
                axes[0].set_xticklabels(env_indicators_short, rotation=45, ha='right')
                axes[0].legend()
                axes[0].set_ylim(0, 100)
                
                # 学生发展指标（动态）
                dev_indicators_short = [self.indicator_names.get(col, col) 
                                        for col in self.dev_indicators]
                x = np.arange(len(dev_indicators_short))
                
                for i, gender in enumerate(self.gender_values[:2]):
                    gender_data = gender_analysis[gender_analysis['性别'] == gender]
                    if len(gender_data) > 0:
                        values = [gender_data.get(short, 0).values[0] if short in gender_data.columns else 0
                                  for short in dev_indicators_short]
                        if i == 0:
                            bars1 = axes[1].bar(x - width/2, values, width, 
                                                label=self.gender_values[0], color='#4472C4')
                        else:
                            bars2 = axes[1].bar(x + width/2, values, width, 
                                                label=self.gender_values[1], color='#ED7D31')
                
                axes[1].set_ylabel('平均分', fontsize=11)
                axes[1].set_title('学生发展维度性别差异', fontsize=12, fontweight='bold')
                axes[1].set_xticks(x)
                axes[1].set_xticklabels(dev_indicators_short, rotation=45, ha='right')
                axes[1].legend()
                axes[1].set_ylim(0, 100)
                
                plt.suptitle('性别差异分析', fontsize=14, fontweight='bold', y=1.02)
                plt.tight_layout()
                plt.savefig(charts_dir / '性别差异分析.png', bbox_inches='tight')
                plt.close()
                
                if verbose:
                    print("图表：性别差异分析 已保存")
        
        # 图4: 学校层面定位图（动态）
        school_position_path = self.output_dir / 'data' / 'school_position.csv'
        if school_position_path.exists():
            school_stats = pd.read_csv(school_position_path)
            
            if len(school_stats) > 0:
                fig, ax = plt.subplots(figsize=(12, 8))
                scatter = ax.scatter(school_stats['成长环境_均分'], school_stats['学生发展_均分'], 
                                     s=school_stats['学生人数']/2, alpha=0.6, c='#4472C4', 
                                     edgecolors='black', linewidth=0.5)
                
                ax.set_xlabel('成长环境均分', fontsize=12)
                ax.set_ylabel('学生发展均分', fontsize=12)
                ax.set_title('学校层面环境-发展定位图（气泡大小代表样本量）', fontsize=14, fontweight='bold')
                
                # 动态中位数线
                ax.axhline(y=school_stats['学生发展_均分'].median(), 
                           color='red', linestyle='--', linewidth=1, alpha=0.5)
                ax.axvline(x=school_stats['成长环境_均分'].median(), 
                           color='red', linestyle='--', linewidth=1, alpha=0.5)
                
                # 标注前5和后5学校
                if len(school_stats) >= 5:
                    top5_schools = school_stats.nlargest(5, '学生发展_均分')['学校代码'].tolist()
                    bottom5_schools = school_stats.nsmallest(5, '学生发展_均分')['学校代码'].tolist()
                    
                    for _, row in school_stats.iterrows():
                        if row['学校代码'] in top5_schools or row['学校代码'] in bottom5_schools:
                            ax.annotate(str(row['学校代码']), 
                                       (row['成长环境_均分'], row['学生发展_均分']), 
                                       fontsize=8, ha='center', va='bottom')
                
                # 动态范围
                ax.set_xlim(school_stats['成长环境_均分'].min() - 5, 
                           school_stats['成长环境_均分'].max() + 5)
                ax.set_ylim(school_stats['学生发展_均分'].min() - 5, 
                           school_stats['学生发展_均分'].max() + 5)
                
                plt.tight_layout()
                plt.savefig(charts_dir / '学校层面定位图.png', bbox_inches='tight')
                plt.close()
                
                if verbose:
                    print("图表：学校层面定位图 已保存")
        
        # 图5: 学生四象限分布（动态）
        quadrant_path = self.output_dir / 'data' / 'quadrant_analysis.csv'
        if quadrant_path.exists():
            quadrant_stats = pd.read_csv(quadrant_path)
            
            if len(quadrant_stats) > 0:
                quadrant_order = ['第一象限-优势发展型', '第二象限-逆境成长型', 
                                 '第四象限-重点关注型', '第三象限-潜力待发型']
                quadrant_data = quadrant_stats[quadrant_stats['象限'].isin(quadrant_order)]
                quadrant_data = quadrant_data.set_index('象限').reindex(quadrant_order).reset_index()
                
                fig, ax = plt.subplots(figsize=(10, 6))
                colors = ['#90EE90', '#87CEEB', '#FFB6C1', '#FF6B6B']
                bars = ax.bar(range(len(quadrant_data)), quadrant_data['人数'].fillna(0), color=colors)
                ax.set_xticks(range(len(quadrant_data)))
                ax.set_xticklabels(['第一象限\n优势发展型', '第二象限\n逆境成长型', 
                                   '第四象限\n重点关注型', '第三象限\n潜力待发型'], fontsize=10)
                ax.set_ylabel('学生人数', fontsize=12)
                ax.set_title('学生环境-发展四象限分布', fontsize=14, fontweight='bold')
                
                total = quadrant_data['人数'].fillna(0).sum()
                if total > 0:
                    for i, (bar, value) in enumerate(zip(bars, quadrant_data['人数'].fillna(0))):
                        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 50, 
                                f'{int(value)}\n({value/total*100:.1f}%)', ha='center', va='bottom', fontsize=10)
                
                plt.tight_layout()
                plt.savefig(charts_dir / '学生四象限分布.png', bbox_inches='tight')
                plt.close()
                
                if verbose:
                    print("图表：学生四象限分布 已保存")
        
        # 图6: 办学规模对比（动态）
        scale_path = self.output_dir / 'data' / 'scale_analysis.csv'
        if scale_path.exists():
            scale_analysis = pd.read_csv(scale_path)
            
            if len(scale_analysis) > 0:
                fig, ax = plt.subplots(figsize=(max(8, len(scale_analysis) * 1.5), 6))
                x = np.arange(len(scale_analysis))
                width = 0.35
                
                bars1 = ax.bar(x - width/2, scale_analysis['成长环境均分'], width, 
                               label='成长环境均分', color='#4472C4')
                bars2 = ax.bar(x + width/2, scale_analysis['学生发展均分'], width, 
                               label='学生发展均分', color='#ED7D31')
                
                ax.set_ylabel('平均分', fontsize=12)
                ax.set_title('不同办学规模的环境与发展均分', fontsize=14, fontweight='bold')
                ax.set_xticks(x)
                ax.set_xticklabels(scale_analysis['办学规模'])
                ax.legend()
                ax.set_ylim(0, 100)
                
                for bar, value in zip(bars1, scale_analysis['成长环境均分']):
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                            f'{value:.1f}', ha='center', va='bottom', fontsize=9)
                for bar, value in zip(bars2, scale_analysis['学生发展均分']):
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                            f'{value:.1f}', ha='center', va='bottom', fontsize=9)
                
                plt.tight_layout()
                plt.savefig(charts_dir / '办学规模对比.png', bbox_inches='tight')
                plt.close()
                
                if verbose:
                    print("图表：办学规模对比 已保存")
        
        # 图7: 公办vs民办对比（动态）
        nature_path = self.output_dir / 'data' / 'nature_analysis.csv'
        if nature_path.exists():
            nature_analysis = pd.read_csv(nature_path)
            
            if len(nature_analysis) > 0:
                fig, ax = plt.subplots(figsize=(max(6, len(nature_analysis) * 2), 6))
                x = np.arange(len(nature_analysis))
                width = 0.35
                
                bars1 = ax.bar(x - width/2, nature_analysis['成长环境均分'], width, 
                               label='成长环境均分', color='#4472C4')
                bars2 = ax.bar(x + width/2, nature_analysis['学生发展均分'], width, 
                               label='学生发展均分', color='#ED7D31')
                
                ax.set_ylabel('平均分', fontsize=12)
                ax.set_title(f'{self.nature_values[0] if len(self.nature_values) > 0 else "办学性质"}学校对比', 
                            fontsize=14, fontweight='bold')
                ax.set_xticks(x)
                ax.set_xticklabels(nature_analysis['办学性质'])
                ax.legend()
                ax.set_ylim(0, 100)
                
                for bar, value in zip(bars1, nature_analysis['成长环境均分']):
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                            f'{value:.1f}', ha='center', va='bottom', fontsize=9)
                for bar, value in zip(bars2, nature_analysis['学生发展均分']):
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                            f'{value:.1f}', ha='center', va='bottom', fontsize=9)
                
                plt.tight_layout()
                plt.savefig(charts_dir / '办学性质对比.png', bbox_inches='tight')
                plt.close()
                
                if verbose:
                    print("图表：办学性质对比 已保存")
        
        if verbose:
            print("\n所有图表已生成完成！")
    
    def generate_report(self, verbose=False, recommendations=None):
        """生成专业格式的Word报告
        
        Args:
            verbose: 是否输出详细信息
            recommendations: 自定义改进建议（可选），格式为：
                {
                    'priority': ['建议1', '建议2', ...],
                    'targeted': ['建议1', '建议2', ...],
                    'longterm': ['建议1', '建议2', ...]
                }
                如果为None，则根据数据动态生成
        """
        if verbose:
            print("\n="*60)
            print("生成专业格式Word报告")
            print("="*60)
        
        # 读取分析数据
        indicator_stats_path = self.output_dir / 'data' / 'indicator_stats.csv'
        if not indicator_stats_path.exists():
            if verbose:
                print("警告：缺少分析数据，无法生成报告")
            return None
        
        indicator_stats = pd.read_csv(indicator_stats_path)
        
        # 加载各类分析数据
        gender_analysis = pd.read_csv(self.output_dir / 'data' / 'gender_analysis.csv') \
            if (self.output_dir / 'data' / 'gender_analysis.csv').exists() else pd.DataFrame()
        scale_analysis = pd.read_csv(self.output_dir / 'data' / 'scale_analysis.csv') \
            if (self.output_dir / 'data' / 'scale_analysis.csv').exists() else pd.DataFrame()
        nature_analysis = pd.read_csv(self.output_dir / 'data' / 'nature_analysis.csv') \
            if (self.output_dir / 'data' / 'nature_analysis.csv').exists() else pd.DataFrame()
        school_stats = pd.read_csv(self.output_dir / 'data' / 'school_position.csv') \
            if (self.output_dir / 'data' / 'school_position.csv').exists() else pd.DataFrame()
        quadrant_stats = pd.read_csv(self.output_dir / 'data' / 'quadrant_analysis.csv') \
            if (self.output_dir / 'data' / 'quadrant_analysis.csv').exists() else pd.DataFrame()
        intervention_suggestions = pd.read_csv(self.output_dir / 'data' / 'intervention_suggestions.csv') \
            if (self.output_dir / 'data' / 'intervention_suggestions.csv').exists() else pd.DataFrame()
        
        df_clean = self.df_clean
        total_students = len(df_clean)
        total_schools = df_clean['学校代码'].nunique() if '学校代码' in df_clean.columns else 0
        
        # 创建格式化报告
        try:
            formatter = create_formatted_report(title='教育学情数据分析报告')
        except:
            formatter = ReportFormatter(Document())
            formatter.add_cover_page('教育学情数据分析报告', '数据分析与干预建议')
        
        # ===== 一、数据概况与分析口径 =====
        formatter.add_section_heading('一、数据概况与分析口径', level=1)
        
        formatter.add_paragraph_with_style(
            f'本次分析基于学生指标得分数据，旨在全面了解学生成长环境和学生发展状况，识别需要重点关注的群体和学校，为教育决策提供数据支持。',
            indent=True
        )
        
        formatter.add_section_heading('1. 数据规模', level=2)
        
        scale_data = [
            ['样本总量', f'{total_students}名学生'],
            ['学校数量', f'{total_schools}所学校' if total_schools > 0 else '未统计'],
            ['分析日期', datetime.datetime.now().strftime('%Y年%m月%d日')]
        ]
        formatter.add_table_with_header(['统计项', '数值'], scale_data, title='数据规模统计')
        
        formatter.add_section_heading('2. 分析指标体系', level=2)
        
        formatter.add_paragraph_with_style(
            f'本次分析采用科学的指标体系，涵盖成长环境和学生发展两大维度：',
            indent=True
        )
        
        formatter.add_bullet_point(f'成长环境指标：共{len(self.env_indicators)}个，包括{", ".join([self.indicator_names.get(col, col) for col in self.env_indicators])}')
        formatter.add_bullet_point(f'学生发展指标：共{len(self.dev_indicators)}个，包括{", ".join([self.indicator_names.get(col, col) for col in self.dev_indicators])}')
        
        formatter.add_paragraph_with_style('', indent=True)
        
        formatter.add_section_heading('3. 样本分布特征', level=2)
        
        if '性别' in df_clean.columns and len(self.gender_values) > 0:
            gender_dist = df_clean['性别'].value_counts()
            gender_data = []
            for gender in self.gender_values:
                count = gender_dist.get(gender, 0)
                gender_data.append([gender, f'{count}人', f'{count/total_students*100:.2f}%'])
            formatter.add_table_with_header(['性别', '人数', '占比'], gender_data)
        
        if '办学性质' in df_clean.columns and len(self.nature_values) > 0:
            nature_dist = df_clean['办学性质'].value_counts()
            nature_data = []
            for nature in self.nature_values:
                count = nature_dist.get(nature, 0)
                nature_data.append([nature, f'{count}人', f'{count/total_students*100:.2f}%'])
            formatter.add_table_with_header(['办学性质', '人数', '占比'], nature_data)
        
        if '办学规模' in df_clean.columns and len(self.scale_values) > 0:
            scale_dist = df_clean['办学规模'].value_counts()
            scale_data = []
            for scale in self.scale_values:
                count = scale_dist.get(scale, 0)
                scale_data.append([scale, f'{count}人', f'{count/total_students*100:.2f}%'])
            formatter.add_table_with_header(['办学规模', '人数', '占比'], scale_data)
        
        formatter.add_page_break()
        
        # ===== 二、整体学情画像 =====
        formatter.add_section_heading('二、整体学情画像分析', level=1)
        
        formatter.add_paragraph_with_style(
            '通过分析各指标的平均分和风险暴露率，可以全面了解学生的整体学情状况，识别优势指标和风险指标。',
            indent=True
        )
        
        # 图: 各指标平均分对比
        if len(indicator_stats) > 0 and (self.output_dir / 'report_charts' / '各指标平均分对比.png').exists():
            formatter.add_figure(
                self.output_dir / 'report_charts' / '各指标平均分对比.png',
                '各指标平均分对比'
            )
            
            # 添加图表描述与教育分析
            if len(indicator_stats) > 0:
                overall_avg = indicator_stats['平均分'].mean()
                max_idx = indicator_stats['平均分'].idxmax()
                min_idx = indicator_stats['平均分'].idxmin()
                
                desc = f'如图所示，{len(indicator_stats)}个指标的平均得分整体为{overall_avg:.2f}分。'
                
                if pd.notna(max_idx):
                    max_indicator = indicator_stats.loc[max_idx]
                    desc += f'其中{max_indicator["指标"]}得分最高，达到{max_indicator["平均分"]:.2f}分，'
                    desc += '表明该维度表现优秀，可作为保持和推广的重点。'
                
                if pd.notna(min_idx):
                    min_indicator = indicator_stats.loc[min_idx]
                    desc += f'{min_indicator["指标"]}得分最低，仅为{min_indicator["平均分"]:.2f}分，'
                    desc += '需要重点关注和改进。'
                
                formatter.add_paragraph_with_style(desc, indent=True)
                
                # 添加教育分析与讨论
                edu_analysis = '【教育分析与讨论】从教育学视角来看，学生发展是学校教育的核心目标。'
                if overall_avg >= 70:
                    edu_analysis += '整体得分处于良好水平，说明学校在学生培养方面取得了积极成效。'
                elif overall_avg >= 60:
                    edu_analysis += '整体得分处于中等水平，仍有较大提升空间，建议从课程设置、教学方法、家校合作等方面进行优化。'
                else:
                    edu_analysis += '整体得分偏低，需要系统性地审视教育过程，从学校管理、教师培训、学生辅导等多维度进行改进。'
                edu_analysis += '建议学校根据指标差异，针对性地制定改进方案，促进学生全面发展。'
                formatter.add_paragraph_with_style(edu_analysis, indent=True)
        
        # 图: 各指标风险暴露率
        if (self.output_dir / 'report_charts' / '各指标风险暴露率.png').exists():
            formatter.add_figure(
                self.output_dir / 'report_charts' / '各指标风险暴露率.png',
                '各指标风险暴露率（低于60分）'
            )
            
            # 添加图表描述与教育分析
            if len(indicator_stats) > 0:
                high_risk = indicator_stats[indicator_stats['风险暴露率(%)'] > 40]
                medium_risk = indicator_stats[(indicator_stats['风险暴露率(%)'] > 20) & (indicator_stats['风险暴露率(%)'] <= 40)]
                low_risk = indicator_stats[indicator_stats['风险暴露率(%)'] <= 20]
                
                desc = f'风险暴露率反映各指标中得分低于60分的学生比例，是识别教育薄弱环节的重要依据。'
                
                if len(high_risk) > 0:
                    high_names = '、'.join([row['指标'] for _, row in high_risk.iterrows()])
                    desc += f'高风险指标（>40%）有{len(high_risk)}个：{high_names}，这些指标亟需开展专项干预。'
                
                if len(medium_risk) > 0:
                    desc += f'中风险指标（20%-40%）有{len(medium_risk)}个，需要建立预警机制。'
                
                if len(low_risk) > 0:
                    desc += f'低风险指标（≤20%）有{len(low_risk)}个，表现良好可继续保持。'
                
                formatter.add_paragraph_with_style(desc, indent=True)
                
                # 添加教育分析与讨论
                edu_analysis = '【教育分析与讨论】风险暴露率揭示了学生群体中需要特别关注的比例。'
                if len(high_risk) > 0:
                    edu_analysis += f'高风险指标的存在说明部分学生在某些维度存在明显短板，'
                    edu_analysis += '这可能与家庭教育环境、学校资源配置、学生个体差异等因素有关。'
                    edu_analysis += '教育工作者应建立"一生一策"的帮扶机制，通过个性化辅导、同伴互助、家校联动等方式，帮助学生突破发展瓶颈。'
                if len(medium_risk) > 0:
                    edu_analysis += '中风险指标提示存在潜在的教育隐患，应建立早期预警和干预机制，防患于未然。'
                formatter.add_paragraph_with_style(edu_analysis, indent=True)
        
        formatter.add_page_break()
        
        # ===== 三、群体差异分析 =====
        formatter.add_section_heading('三、群体差异分析', level=1)
        
        formatter.add_paragraph_with_style(
            '不同群体学生在成长环境和学生发展方面存在差异，了解这些差异有助于制定针对性的干预措施。',
            indent=True
        )
        
        if len(gender_analysis) > 0 and (self.output_dir / 'report_charts' / '性别差异分析.png').exists():
            formatter.add_section_heading('3.1 性别差异分析', level=2)
            formatter.add_figure(
                self.output_dir / 'report_charts' / '性别差异分析.png',
                '性别差异对比'
            )
            
            if len(gender_analysis) >= 2:
                gender_table_data = []
                for _, row in gender_analysis.iterrows():
                    gender_table_data.append([
                        row['性别'],
                        int(row['人数']),
                        f'{row["成长环境均分"]:.2f}',
                        f'{row["学生发展均分"]:.2f}'
                    ])
                formatter.add_table_with_header(
                    ['性别', '人数', '成长环境均分', '学生发展均分'],
                    gender_table_data,
                    title='性别差异统计'
                )
                
                # 添加表格描述与教育分析
                male_data = gender_analysis[gender_analysis['性别'].isin(['男', '男生'])]
                female_data = gender_analysis[gender_analysis['性别'].isin(['女', '女生'])]
                
                if len(male_data) > 0 and len(female_data) > 0:
                    env_gap = abs(male_data['成长环境均分'].values[0] - female_data['成长环境均分'].values[0])
                    dev_gap = abs(male_data['学生发展均分'].values[0] - female_data['学生发展均分'].values[0])
                    
                    desc = f'统计显示，男生{int(male_data["人数"].values[0])}人，女生{int(female_data["人数"].values[0])}人。'
                    if env_gap > 2 or dev_gap > 2:
                        desc += f'性别差异较为明显，成长环境均分差距{env_gap:.2f}分，学生发展均分差距{dev_gap:.2f}分。'
                    else:
                        desc += f'性别差异不大，成长环境均分差距仅{env_gap:.2f}分，学生发展均分差距仅{dev_gap:.2f}分。'
                    
                    formatter.add_paragraph_with_style(desc, indent=True)
                    
                    # 添加教育分析与讨论
                    edu_analysis = '【教育分析与讨论】性别差异研究是教育公平的重要议题。'
                    if env_gap > 2 or dev_gap > 2:
                        edu_analysis += '数据显示男女生在某些维度存在明显差异，'
                        edu_analysis += '这可能与生理发展特点、家庭教育方式、学校教育模式等因素相关。'
                        edu_analysis += '教育者应在承认性别差异的基础上，为不同性别学生提供适切的教育支持，'
                        edu_analysis += '既要避免性别刻板印象的限制，也要关注不同群体的特殊需求，实现因材施教。'
                    else:
                        edu_analysis += '性别差异不明显，说明学校在促进性别平等方面取得了良好成效。'
                        edu_analysis += '建议继续保持公平的教育环境，让每位学生都能获得平等的发展机会。'
                    formatter.add_paragraph_with_style(edu_analysis, indent=True)
        
        if len(scale_analysis) > 0 and (self.output_dir / 'report_charts' / '办学规模对比.png').exists():
            formatter.add_section_heading('3.2 办学规模差异分析', level=2)
            formatter.add_figure(
                self.output_dir / 'report_charts' / '办学规模对比.png',
                '不同办学规模的环境与发展均分'
            )
            
            scale_table_data = []
            for _, row in scale_analysis.iterrows():
                scale_table_data.append([
                    row['办学规模'],
                    int(row['人数']),
                    f'{row["成长环境均分"]:.2f}',
                    f'{row["学生发展均分"]:.2f}'
                ])
            formatter.add_table_with_header(
                ['办学规模', '人数', '成长环境均分', '学生发展均分'],
                scale_table_data,
                title='办学规模差异统计'
            )
            
            # 添加表格描述与教育分析
            if len(scale_analysis) > 1:
                max_dev = scale_analysis.loc[scale_analysis['学生发展均分'].idxmax()]
                min_dev = scale_analysis.loc[scale_analysis['学生发展均分'].idxmin()]
                gap = max_dev['学生发展均分'] - min_dev['学生发展均分']
                
                desc = f'不同办学规模学校学生表现存在差异。'
                desc += f'{max_dev["办学规模"]}学校学生发展均分最高（{max_dev["学生发展均分"]:.2f}分），'
                desc += f'{min_dev["办学规模"]}学校最低（{min_dev["学生发展均分"]:.2f}分），'
                desc += f'差距为{gap:.2f}分。'
                
                formatter.add_paragraph_with_style(desc, indent=True)
                
                # 添加教育分析与讨论
                edu_analysis = '【教育分析与讨论】办学规模是影响教育质量的重要因素。'
                if gap > 5:
                    edu_analysis += f'不同规模学校学生发展差距达到{gap:.2f}分，差异较为明显。'
                    edu_analysis += '大规模学校可能拥有更丰富的教育资源和更完善的课程体系，'
                    edu_analysis += '而小规模学校则可能面临师资不足、课程开设受限等问题。'
                    edu_analysis += '教育主管部门应优化资源配置，对小规模学校给予更多支持，'
                    edu_analysis += '如加强师资培训、推进集团化办学、共享优质教育资源等，缩小校际差距。'
                else:
                    edu_analysis += '不同规模学校差距不大，说明教育资源配置相对均衡。'
                    edu_analysis += '建议继续优化资源配置，促进各规模学校协调发展。'
                formatter.add_paragraph_with_style(edu_analysis, indent=True)
        
        if len(nature_analysis) > 0 and (self.output_dir / 'report_charts' / '办学性质对比.png').exists():
            formatter.add_section_heading('3.3 办学性质差异分析', level=2)
            formatter.add_figure(
                self.output_dir / 'report_charts' / '办学性质对比.png',
                '不同办学性质学校对比'
            )
            
            nature_table_data = []
            for _, row in nature_analysis.iterrows():
                nature_table_data.append([
                    row['办学性质'],
                    int(row['人数']),
                    f'{row["成长环境均分"]:.2f}',
                    f'{row["学生发展均分"]:.2f}'
                ])
            formatter.add_table_with_header(
                ['办学性质', '人数', '成长环境均分', '学生发展均分'],
                nature_table_data,
                title='办学性质差异统计'
            )
            
            # 添加表格描述与教育分析
            if len(nature_analysis) > 1:
                desc = f'不同办学性质学校学生表现对比：'
                for _, row in nature_analysis.iterrows():
                    desc += f'{row["办学性质"]}学校{int(row["人数"])}人，学生发展均分{row["学生发展均分"]:.2f}分；'
                formatter.add_paragraph_with_style(desc, indent=True)
                
                # 添加教育分析与讨论
                edu_analysis = '【教育分析与讨论】公办与民办学校各有特点和优势。'
                edu_analysis += '公办学校承担着基础教育的主要职责，民办学校则为家长提供了多元化的教育选择。'
                edu_analysis += '无论办学性质如何，都应以促进学生全面发展为目标，不断优化教育质量。'
                edu_analysis += '教育主管部门应建立公平竞争、相互促进的制度环境，推动各类学校共同发展。'
                formatter.add_paragraph_with_style(edu_analysis, indent=True)
        
        formatter.add_page_break()
        
        # ===== 四、学校层面分析 =====
        if len(school_stats) > 0 and (self.output_dir / 'report_charts' / '学校层面定位图.png').exists():
            formatter.add_section_heading('四、学校层面定位分析', level=1)
            
            formatter.add_paragraph_with_style(
                '学校层面分析通过定位图展示各学校在成长环境和学生发展两个维度上的相对位置，帮助识别表现优秀和需要帮扶的学校。',
                indent=True
            )
            
            formatter.add_figure(
                self.output_dir / 'report_charts' / '学校层面定位图.png',
                '学校环境-发展定位图（气泡大小代表样本量）'
            )
            
            # 添加图表描述与教育分析
            if len(school_stats) > 0:
                avg_dev = school_stats['学生发展_均分'].mean()
                std_dev = school_stats['学生发展_均分'].std()
                desc = f'定位图展示了{len(school_stats)}所学校的学生发展状况。'
                desc += f'各校学生发展均分平均为{avg_dev:.2f}分，标准差为{std_dev:.2f}分。'
                formatter.add_paragraph_with_style(desc, indent=True)
                
                # 添加教育分析与讨论
                edu_analysis = '【教育分析与讨论】校际差异是教育均衡发展的重要观测指标。'
                if std_dev > 5:
                    edu_analysis += f'学校间标准差达到{std_dev:.2f}分，差异较大，'
                    edu_analysis += '反映出教育资源配置、学校管理水平、师资力量等方面存在不平衡。'
                    edu_analysis += '教育行政部门应建立校际帮扶机制，推广优秀学校的办学经验，'
                    edu_analysis += '对薄弱学校进行精准帮扶，缩小校际差距，促进教育公平。'
                else:
                    edu_analysis += '学校间差异相对较小，说明区域内教育发展较为均衡。'
                    edu_analysis += '建议继续保持均衡发展态势，同时鼓励各校形成特色，实现"各美其美"的发展格局。'
                formatter.add_paragraph_with_style(edu_analysis, indent=True)
            
            if len(school_stats) >= 5:
                formatter.add_section_heading('4.1 表现优秀学校（前5名）', level=2)
                top5 = school_stats.nlargest(5, '学生发展_均分')
                top5_data = []
                for _, row in top5.iterrows():
                    top5_data.append([
                        str(row['学校代码']),
                        int(row['学生人数']),
                        f'{row["成长环境_均分"]:.2f}',
                        f'{row["学生发展_均分"]:.2f}'
                    ])
                formatter.add_table_with_header(
                    ['学校代码', '学生人数', '成长环境均分', '学生发展均分'],
                    top5_data,
                    title='表现优秀学校统计'
                )
                
                # 添加表格描述
                top5_avg = top5['学生发展_均分'].mean()
                desc = f'表现优秀的5所学校学生发展均分平均为{top5_avg:.2f}分，'
                desc += '可作为标杆学校进行经验总结和推广。'
                formatter.add_paragraph_with_style(desc, indent=True)
                
                formatter.add_section_heading('4.2 表现欠佳学校（后5名）', level=2)
                bottom5 = school_stats.nsmallest(5, '学生发展_均分')
                bottom5_data = []
                for _, row in bottom5.iterrows():
                    bottom5_data.append([
                        str(row['学校代码']),
                        int(row['学生人数']),
                        f'{row["成长环境_均分"]:.2f}',
                        f'{row["学生发展_均分"]:.2f}'
                    ])
                formatter.add_table_with_header(
                    ['学校代码', '学生人数', '成长环境均分', '学生发展均分'],
                    bottom5_data,
                    title='表现欠佳学校统计'
                )
                
                # 添加表格描述
                bottom5_avg = bottom5['学生发展_均分'].mean()
                gap = top5_avg - bottom5_avg
                desc = f'表现欠佳的5所学校学生发展均分平均为{bottom5_avg:.2f}分。'
                desc += f'优秀学校与欠佳学校的差距为{gap:.2f}分，'
                if gap > 10:
                    desc += '差距较大，需要对这些学校进行重点帮扶。'
                formatter.add_paragraph_with_style(desc, indent=True)
            
            formatter.add_page_break()
        
        # ===== 五、学生分层分析 =====
        if len(quadrant_stats) > 0 and (self.output_dir / 'report_charts' / '学生四象限分布.png').exists():
            formatter.add_section_heading('五、学生分层与干预优先级', level=1)
            
            formatter.add_paragraph_with_style(
                '基于成长环境和学生发展两个维度的中位数，将学生划分为四个象限，为精准干预提供依据。',
                indent=True
            )
            
            formatter.add_figure(
                self.output_dir / 'report_charts' / '学生四象限分布.png',
                '学生环境-发展四象限分布'
            )
            
            # 添加图表描述与教育分析
            if len(quadrant_stats) > 0:
                q1 = quadrant_stats[quadrant_stats['象限'] == '第一象限-优势发展型']
                q4 = quadrant_stats[quadrant_stats['象限'] == '第四象限-重点关注型']
                
                desc = '四象限分布图清晰展示了学生在成长环境与学生发展两个维度上的分布情况。'
                
                if len(q1) > 0:
                    q1_count = int(q1['人数'].values[0])
                    q1_pct = q1_count / total_students * 100
                    desc += f'第一象限（优势发展型）学生{q1_count}人，占比{q1_pct:.1f}%，表现优秀应继续保持。'
                
                if len(q4) > 0:
                    q4_count = int(q4['人数'].values[0])
                    q4_pct = q4_count / total_students * 100
                    desc += f'第四象限（重点关注型）学生{q4_count}人，占比{q4_pct:.1f}%，需要优先干预。'
                
                formatter.add_paragraph_with_style(desc, indent=True)
                
                # 添加教育分析与讨论
                edu_analysis = '【教育分析与讨论】四象限分层体现了生态发展理论的核心观点：'
                edu_analysis += '学生发展是个体与环境互动的结果。'
                edu_analysis += '第一象限学生处于良好的发展生态中，应继续优化教育环境，促进其向更高水平发展。'
                edu_analysis += '第二象限学生展现出较强的心理韧性，但环境支持不足，应改善其成长环境，巩固发展成果。'
                edu_analysis += '第三象限学生环境条件较好但发展滞后，需要深入了解其内在需求，激发发展动力。'
                edu_analysis += '第四象限学生处于双重不利境地，需要学校、家庭、社会多方协同，建立综合帮扶体系。'
                edu_analysis += '教育干预应遵循"预防为先、干预为重"的原则，实施差异化、精准化的教育支持。'
                formatter.add_paragraph_with_style(edu_analysis, indent=True)
            
            formatter.add_section_heading('5.1 四象限划分标准', level=2)
            formatter.add_bullet_point('第一象限（优势发展型）：成长环境≥中位数 且 学生发展≥中位数')
            formatter.add_bullet_point('第二象限（逆境成长型）：成长环境<中位数 且 学生发展≥中位数')
            formatter.add_bullet_point('第三象限（潜力待发型）：成长环境≥中位数 且 学生发展<中位数')
            formatter.add_bullet_point('第四象限（重点关注型）：成长环境<中位数 且 学生发展<中位数')
            
            formatter.add_section_heading('5.2 各象限学生分布', level=2)
            
            if len(quadrant_stats) > 0:
                quadrant_table_data = []
                quadrant_order = ['第一象限-优势发展型', '第二象限-逆境成长型', 
                                 '第三象限-潜力待发型', '第四象限-重点关注型']
                
                for quadrant in quadrant_order:
                    row_data = quadrant_stats[quadrant_stats['象限'] == quadrant]
                    if len(row_data) > 0:
                        count = int(row_data['人数'].values[0])
                        quadrant_table_data.append([
                            quadrant.split('-')[0],
                            quadrant.split('-')[1] if '-' in quadrant else quadrant,
                            count,
                            f'{count/total_students*100:.2f}%'
                        ])
                
                formatter.add_table_with_header(
                    ['象限', '类型', '人数', '占比'],
                    quadrant_table_data,
                    title='四象限学生分布统计'
                )
                
                # 添加表格描述
                desc = '四象限分层为精准干预提供了依据：第一象限学生可继续保持优势；'
                desc += '第二象限学生需要关注环境因素，预防潜在风险；'
                desc += '第三象限学生需要激发发展潜力；'
                desc += '第四象限学生需要优先干预，建立多维度帮扶体系。'
                formatter.add_paragraph_with_style(desc, indent=True)
            
            formatter.add_section_heading('5.3 干预建议', level=2)
            
            if len(intervention_suggestions) > 0:
                for _, row in intervention_suggestions.iterrows():
                    formatter.add_paragraph_with_style(
                        f'{row["干预建议"]}',
                        bold_prefix=f'{row["象限"]}（{int(row["人数"])}人）：',
                        indent=True
                    )
            
            formatter.add_page_break()
        
        # ===== 六、结论与建议 =====
        formatter.add_section_heading('六、结论与改进建议', level=1)
        
        # 6.1 主要结论 - 根据数据动态生成
        formatter.add_section_heading('6.1 主要结论', level=2)
        
        conclusions = self._generate_conclusions(
            indicator_stats, gender_analysis, scale_analysis, nature_analysis,
            school_stats, quadrant_stats, total_students, total_schools
        )
        formatter.add_analysis_summary(conclusions)
        
        # 6.2 改进建议 - 根据数据动态生成或使用自定义建议
        formatter.add_section_heading('6.2 改进建议', level=2)
        
        if recommendations is None:
            recommendations = self._generate_recommendations(
                indicator_stats, gender_analysis, scale_analysis, nature_analysis,
                school_stats, quadrant_stats
            )
        
        formatter.add_recommendations(recommendations)
        
        # 保存文档
        report_path = self.output_dir / 'report' / '教育学情数据分析报告.docx'
        formatter.save(report_path)
        
        if verbose:
            print(f"专业格式Word报告已生成：{report_path}")
        
        return report_path
    
    def _generate_conclusions(self, indicator_stats, gender_analysis, scale_analysis,
                              nature_analysis, school_stats, quadrant_stats,
                              total_students, total_schools):
        """根据分析数据动态生成主要结论"""
        conclusions = []
        
        # 结论1：数据规模
        conclusion1 = f'本次分析覆盖{total_students}名学生'
        if total_schools > 0:
            conclusion1 += f'，{total_schools}所学校'
        conclusion1 += f'，共分析{len(self.env_indicators)}个成长环境指标和{len(self.dev_indicators)}个学生发展指标。'
        conclusions.append(conclusion1)
        
        # 结论2：整体学情
        if len(indicator_stats) > 0:
            overall_avg = indicator_stats['平均分'].mean()
            high_risk_count = len(indicator_stats[indicator_stats['风险暴露率(%)'] > 40])
            
            if overall_avg >= 70:
                conclusions.append(f'整体学情表现良好，各指标平均得分为{overall_avg:.2f}分。')
            elif overall_avg >= 60:
                conclusions.append(f'整体学情表现中等，各指标平均得分为{overall_avg:.2f}分，仍有提升空间。')
            else:
                conclusions.append(f'整体学情表现欠佳，各指标平均得分为{overall_avg:.2f}分，需要重点关注。')
            
            if high_risk_count > 0:
                conclusions.append(f'共有{high_risk_count}个高风险指标（风险暴露率>40%），需要优先干预。')
        
        # 结论3：群体差异
        if len(gender_analysis) >= 2:
            male_data = gender_analysis[gender_analysis['性别'] == gender_analysis['性别'].iloc[0]]
            female_data = gender_analysis[gender_analysis['性别'] == gender_analysis['性别'].iloc[1]]
            if len(male_data) > 0 and len(female_data) > 0:
                gender_gap = abs(male_data['学生发展均分'].values[0] - female_data['学生发展均分'].values[0])
                if gender_gap > 2:
                    conclusions.append(f'性别差异明显，不同性别学生发展均分差距为{gender_gap:.2f}分。')
        
        if len(school_stats) >= 5:
            top5 = school_stats.nlargest(5, '学生发展_均分')
            bottom5 = school_stats.nsmallest(5, '学生发展_均分')
            school_gap = top5['学生发展_均分'].mean() - bottom5['学生发展_均分'].mean()
            conclusions.append(f'学校层面差异显著，表现优秀与欠佳学校的发展均分差距为{school_gap:.2f}分。')
        
        # 结论4：学生分层
        if len(quadrant_stats) > 0:
            q4 = quadrant_stats[quadrant_stats['象限'] == '第四象限-重点关注型']
            if len(q4) > 0:
                q4_count = int(q4['人数'].values[0])
                q4_ratio = q4_count / total_students * 100
                conclusions.append(f'学生分层方面，第四象限（重点关注型）学生共{q4_count}人（{q4_ratio:.1f}%），需要优先关注。')
        
        return conclusions
    
    def _generate_recommendations(self, indicator_stats, gender_analysis, scale_analysis,
                                   nature_analysis, school_stats, quadrant_stats):
        """根据分析数据动态生成改进建议
        
        返回格式：[('类别标题', ['建议1', '建议2', ...]), ...]
        """
        priority_items = []
        targeted_items = []
        longterm_items = []
        
        # 根据高风险指标生成优先干预建议
        if len(indicator_stats) > 0:
            high_risk = indicator_stats[indicator_stats['风险暴露率(%)'] > 40]
            if len(high_risk) > 0:
                high_risk_names = '、'.join([row['指标'] for _, row in high_risk.iterrows()])
                priority_items.append(f'针对高风险指标（{high_risk_names}）开展专项干预项目')
            
            medium_risk = indicator_stats[(indicator_stats['风险暴露率(%)'] > 20) & (indicator_stats['风险暴露率(%)'] <= 40)]
            if len(medium_risk) > 0:
                priority_items.append('对中风险指标建立预警监测机制')
        
        # 根据四象限分析生成建议
        if len(quadrant_stats) > 0:
            q4 = quadrant_stats[quadrant_stats['象限'] == '第四象限-重点关注型']
            if len(q4) > 0:
                priority_items.append('建立第四象限学生多维度帮扶体系，制定个性化发展计划')
            
            q2 = quadrant_stats[quadrant_stats['象限'] == '第二象限-逆境成长型']
            if len(q2) > 0:
                priority_items.append('关注第二象限学生的心理韧性，预防环境因素导致的潜在风险')
        
        # 根据学校分析生成针对性建议
        if len(school_stats) >= 5:
            bottom5 = school_stats.nsmallest(5, '学生发展_均分')
            targeted_items.append(f'对表现欠佳的5所学校进行重点帮扶')
            targeted_items.append('加强学校间资源共享与交流合作')
        
        # 根据群体差异生成建议
        if len(gender_analysis) >= 2:
            targeted_items.append('关注性别差异，提供差异化辅导方案')
        
        if len(scale_analysis) > 1:
            targeted_items.append('针对不同办学规模学校制定差异化支持策略')
        
        # 长效机制建议（通用）
        longterm_items.append('建立定期评估与监测机制')
        longterm_items.append('建立预警机制，及时发现潜在问题')
        longterm_items.append('建立资源调配机制，优化教育资源配置')
        longterm_items.append('建立家校协同机制，形成教育合力')
        
        # 确保每个类别至少有2条建议
        if len(priority_items) < 2:
            priority_items.append('加强心理健康教育和同伴关系建设')
        if len(targeted_items) < 2:
            targeted_items.append('建立教师专业发展支持体系')
        
        return [
            ('优先干预措施', priority_items),
            ('针对性改进措施', targeted_items),
            ('长效机制建设', longterm_items)
        ]
    
    def get_analysis_summary(self):
        """获取分析结果摘要（供大模型重新编写建议时参考）
        
        返回包含关键分析结果的字典
        """
        if self.df_clean is None:
            return None
        
        df_clean = self.df_clean
        total_students = len(df_clean)
        total_schools = df_clean['学校代码'].nunique() if '学校代码' in df_clean.columns else 0
        
        summary = {
            'basic_info': {
                'total_students': total_students,
                'total_schools': total_schools,
                'env_indicators': len(self.env_indicators),
                'dev_indicators': len(self.dev_indicators),
                'env_indicator_names': [self.indicator_names.get(col, col) for col in self.env_indicators],
                'dev_indicator_names': [self.indicator_names.get(col, col) for col in self.dev_indicators]
            }
        }
        
        # 读取分析结果
        indicator_stats_path = self.output_dir / 'data' / 'indicator_stats.csv'
        if indicator_stats_path.exists():
            indicator_stats = pd.read_csv(indicator_stats_path)
            summary['indicator_analysis'] = {
                'overall_avg': float(indicator_stats['平均分'].mean()),
                'max_indicator': indicator_stats.loc[indicator_stats['平均分'].idxmax()].to_dict() if len(indicator_stats) > 0 else None,
                'min_indicator': indicator_stats.loc[indicator_stats['平均分'].idxmin()].to_dict() if len(indicator_stats) > 0 else None,
                'high_risk_count': int(len(indicator_stats[indicator_stats['风险暴露率(%)'] > 40])),
                'high_risk_indicators': indicator_stats[indicator_stats['风险暴露率(%)'] > 40]['指标'].tolist()
            }
        
        quadrant_path = self.output_dir / 'data' / 'quadrant_analysis.csv'
        if quadrant_path.exists():
            quadrant_stats = pd.read_csv(quadrant_path)
            summary['quadrant_analysis'] = {}
            for _, row in quadrant_stats.iterrows():
                summary['quadrant_analysis'][row['象限']] = {
                    'count': int(row['人数']),
                    'env_avg': float(row['成长环境_均分']),
                    'dev_avg': float(row['学生发展_均分'])
                }
        
        school_path = self.output_dir / 'data' / 'school_position.csv'
        if school_path.exists():
            school_stats = pd.read_csv(school_path)
            if len(school_stats) >= 5:
                top5 = school_stats.nlargest(5, '学生发展_均分')
                bottom5 = school_stats.nsmallest(5, '学生发展_均分')
                summary['school_analysis'] = {
                    'top5_avg': float(top5['学生发展_均分'].mean()),
                    'bottom5_avg': float(bottom5['学生发展_均分'].mean()),
                    'gap': float(top5['学生发展_均分'].mean() - bottom5['学生发展_均分'].mean())
                }
        
        return summary
    
    def run_analysis(self, verbose=False):
        """执行完整分析流程"""
        # 1. 加载数据
        self.load_data(verbose)
        
        # 2. 整体学情分析
        self.analyze_overall(verbose)
        
        # 3. 群体差异分析
        self.analyze_group_differences(verbose)
        
        # 4. 学校层面分析
        self.analyze_school_level(verbose)
        
        # 5. 四象限分析
        self.analyze_quadrant(verbose)
        
        # 6. 生成图表
        self.generate_charts(verbose)
        
        # 7. 生成报告
        report_path = self.generate_report(verbose)
        
        if verbose:
            print("\n="*60)
            print("分析完成！")
            print("="*60)
            print(f"\n输出目录：{self.output_dir}")
            print(f"Word报告：{report_path}")
            print(f"数据文件：{self.output_dir / 'data'}")
            print(f"图表文件：{self.output_dir / 'report_charts'}")
            print(f"\n数据配置已保存至：{self.output_dir / 'data' / 'data_config.json'}")
        
        return report_path


def parse_args():
    """解析命令行参数"""
    parser = argparse.ArgumentParser(description='教育学情数据分析脚本（动态版本）')
    parser.add_argument('--data', required=True, help='数据文件路径')
    parser.add_argument('--output', default='output', help='输出目录')
    parser.add_argument('--env-prefix', default='成长环境_', help='成长环境指标前缀')
    parser.add_argument('--dev-prefix', default='学生发展_', help='学生发展指标前缀')
    parser.add_argument('--verbose', action='store_true', help='详细输出模式')
    return parser.parse_args()


def main():
    """主函数"""
    args = parse_args()
    
    # 验证数据文件
    data_path = Path(args.data)
    if not data_path.exists():
        print(f"错误: 数据文件不存在: {data_path}")
        sys.exit(1)
    
    # 创建分析器并执行分析
    analyzer = EducationDataAnalyzer(
        str(data_path), 
        args.output,
        env_prefix=args.env_prefix,
        dev_prefix=args.dev_prefix
    )
    report_path = analyzer.run_analysis(args.verbose)
    
    print(f"\n教育学情数据分析完成！")
    print(f"报告位置: {report_path}")


if __name__ == '__main__':
    main()